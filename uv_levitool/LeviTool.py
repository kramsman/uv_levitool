"""
LeviTool.py creates the attachments need for postcard writing namely scripts and labels.  It does this by performing
a "mail merge" on input files, replacing fields in {} with those from an Excel input file.
All the input files must be in a directory named 'Input', and must contain one xlsx with the info to be merged and
any number of docx files which will have the dat merged into them.
A pdf of each 'docx' is created after the merge.
A separate subdirectory under a min directory named "Output" is created for each used row in the input xlsx file.

To run this in Pycharm, use a configuations set to 'module' with a name like "uv_levitool.LeviTool"
and a working dir like "/Users/Denise/Library/CloudStorage/Dropbox/non ROV python progs/uv_levitool"
"""

# # After creating your project with uv init to create requirements.txt
# uv pip compile pyproject.toml -o requirements.txt

# NEED TO USE BACK VERSION OF OPENPYXL for bug caused in 3.1: county_sheet_df = pd.ExcelFile( script_file_name).parse(
# use openpyxl 3.0.10 not 3.1.2 because later bombs if xlsx has filters  https://stackoverflow.com/questions/75382340/python-pandas-read-excel-error-value-must-be-either-numerical-or-a-string-conta

# TODO currently blanks replaced with "' '" in df used_counties_df.  OK or should be blank? Or '.'?
# TODO Check max length of fields not working. But it's in another version - levitool in non ROV Python?
# TODO: can we grag point size for line in template and compare that ro the allowed (build a table of values and
#  compare)?

# run gitupdater to make sure bekutils and bekgoogle utility libraries are updated
import sys
import os
sys.path.append(os.path.expanduser("~/Dropbox/Postcard Files/"))
if True:
   import gitupdater

import glob
import re
from collections import Counter
from copy import deepcopy
from pathlib import Path

import pandas as pd
from uvbekutils import exit_yes, exit_yes_no, setup_loguru, select_file
from uvbekutils import pyautobek
from check_boe_urls import check_boe_urls, check_short_boe_urls
from docx import Document  # package in Conda is python-docx, not simply docx
from docx2pdf import convert
from loguru import logger
from find_longest_value_in_label_rows import max_label_lengths

pd.options.mode.copy_on_write = True  # fix chain assignment forced in Pandas 3.0

INITIAL_ATTACHMENT_DIR = Path("~/Dropbox/Postcard Files/Attachments/Campaigns/").expanduser()

# TODO can do away specifying length fields and instead use 'keys' list
# Fields to report max length for checking label fit
LENGTH_CHECK_FIELDS = ['CNTYFILENAME', 'PHONE', 'URL']
LENGTH_CHECK_FIELDS = []  #FIXME: Above fields are not what are being used for WI.  Don't the keys in doc work?
# field to use when checking label fit for only rows we are writing
LENGTH_CHECK_FIELDS_SELECT = 'PRIORITY'

setup_loguru("DEBUG", "DEBUG", )

def count_brackets(s):
    """ counts the difference in the number of left and right brackets. returns -10000 if difference is more than 1 """

    bracket_count = s.count('{') - s.count('}')
    if bracket_count < -1 or bracket_count > 1:
        bracket_count = -10000

    return bracket_count


def find_keys(document):
    """ Levi code.  Finds keys three levels deep in document levels."""
    # Keys are fields in document to be replaced
    keys = set()

    d = deepcopy(document)

    for section in d.sections:
        for paragraph in section.header.paragraphs:
            new_keysx = re.findall(r'{[a-zA-Z0-9]+}', paragraph.text)
            new_keys = [x.upper() for x in new_keysx]
            keys = keys.union({k.upper() for k in new_keys})

    for paragraph in d.paragraphs:
        new_keysx = re.findall(r'{[a-zA-Z0-9]+}', paragraph.text)
        new_keys = [x.upper() for x in new_keysx]
        keys = keys.union({k.upper() for k in new_keys})

    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    new_keysx = re.findall(r'{[a-zA-Z0-9]+}', paragraph.text)
                    new_keys = [x.upper() for x in new_keysx]
                    keys = keys.union({k.upper() for k in new_keys})

    return keys


def replace_tags(paragraph, county_row):
    """ Levi code.  Replace tags/keys in paragraph. Runs are weird docx objects. """

    brackets = 0
    partial_runs = []
    for run in paragraph.runs:
        partial_runs.append(run)
        brackets += count_brackets(run.text)
        if brackets == 0:
            text = ''.join([r.text for r in partial_runs])
            for key, val in county_row.items():
                text = re.sub(key, val, text, flags=re.IGNORECASE)  # use re to replace and ignore case
            partial_runs[0].text = text
            for r in partial_runs[1:]:
                r.text = ''

            brackets = 0
            partial_runs = []
        elif brackets != 1:
            msg = f'Template error near "{run.text}" in script template. Make sure sequence to be replaced has consistent formatting.'
            logger.error(msg)
            raise ValueError(msg)


def make_replacements(document, county_row):
    """ Levi code. Call the make replacement code multiple times. """

    for paragraph in document.paragraphs:
        replace_tags(paragraph, county_row)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_tags(paragraph, county_row)

    return


def main_levi():
    """ Read inputs, make replacements.  Easy peasy! """

    def check_header_count(hdr_str, header_count, num_allowed):
        """ make sure headers appear in file only given number of times and error if not"""

        if header_count != num_allowed:
            msg = f"'{hdr_str}' does not appear {num_allowed} time(s), instead {header_count}\n\nEXITING."
            logger.error(msg)
            exit_yes(msg, "ERROR")

    def prompt_for_df_values(df, msg1, box_title):
        """ display values of a dataframe after converting to string and prompt to continue"""

        df_string = df.to_string(index=False)
        logger.debug("here")

        exit_yes_no(f"{msg1}\n\n\n{df_string}\n\n", box_title)

    if True:
        path_input = select_file(
            title="Pick 'Input' Directory",
            start_dir=INITIAL_ATTACHMENT_DIR,
            files_like="Input",
            choices=["Select", "Cancel"],
            mode="dir",  # file, dir or both
            title2="Select the firectory which contains all your template files (docs and the xlsx). Must be named "
                   "'Input'"
        )
        print(f"Selected: {path_input}")
        path_input = Path(path_input)
    else:
        path_input = Path("~/Dropbox/Postcard Files/Attachments/Campaigns/Test box and image/Input").expanduser()
    path_root = path_input.parents[0]

    if path_input.parts[-1].upper() != "INPUT":
        msg = f"Last part of chosen directory:\n\n {path_input}\n\nDoes not equal 'INPUT'.\n\nEXITING."
        logger.error(msg)
        exit_yes(msg, "ERROR")


    # Remove choice of subdirectory for each county since Sincere expects them all in ome dir
    # choice = pyautobek.confirm("Put output files in single Output directory rather than in separate dirs by county ("
    #                            "old way)?",
    #                   title='Single County Dir?',
    #                   buttons=['Yes', 'No', 'Exit'])
    choice = 'yes'
    if choice == 'yes':
        SINGLE_COUNTY_OUTPUT_DIR = True
    elif choice == 'no':
        SINGLE_COUNTY_OUTPUT_DIR = False
    else:
        exit()

    # list of all xlsx files not starting with ~ (temporary files)
    xlsx_files = glob.glob(str(path_input / '[!~]*.xlsx'))

    if len(xlsx_files) != 1:  # must have 1, not more
        msg = (f"Must have only 1 xlsx file in the Input directory; there are {len(xlsx_files)}.  \n\n"
               f"{path_input}\n\n"
               f"EXITING.")
        logger.error(msg)
        exit_yes(msg, "ERROR")
    xlsx_file_name = xlsx_files[0]  # filename of the first (and only) xlsx

    # Get list of all template files (docx) that we will update not starting with ~ (temporary files).
    docx_files = glob.glob(str(path_input / '[!~]*.docx'))
    if not docx_files:
        msg = f"No docx input file found in directory: '{path_input}' EXITING."
        logger.error(msg)
        exit_yes(msg, "ERROR")

    logger.info(f"Will replace tags in {len(docx_files)} documents per county.\n")

    # loop through each docx template file and put keys in required_keys set
    required_keys = set()
    for docx_file in docx_files:
        document_of_docx_file = Document(docx_file)
        keys_in_doc = find_keys(document_of_docx_file)
        required_keys = required_keys.union(keys_in_doc)

    # Add the keys required by the program but not necessarily used for replacement in docxs.
    required_keys.add('{STATE}')
    required_keys.add('{CNTYFILENAME}')  # added BEK 4/23/22
    required_keys.add('{USE}')
    for field in LENGTH_CHECK_FIELDS:
        # add keys for those we are checking max length
        required_keys.add(f"{{{field}}}")
    if LENGTH_CHECK_FIELDS_SELECT:
        # add field used in filtering rows for max length check
        required_keys.add(f"{{{LENGTH_CHECK_FIELDS_SELECT}}}")

    # Note .astype(str) below which reads all data as string and converts missing to 'nan'
    county_sheet_df = pd.read_excel(xlsx_file_name, sheet_name='Counties', header=None, skiprows=1,
                                    nrows=1).astype(str)
    # read only row 2 (1 is skipped above) to get keys
    county_sheet_keys = county_sheet_df.loc[0, :].values.tolist()
    county_sheet_keys = [x.upper() for x in county_sheet_keys if x != "nan"]

    # check that three key tokens only appear once (needed and may not be on file)
    check_header_count('{USE}', county_sheet_keys.count('{USE}'), 1)
    check_header_count('{STATE}', county_sheet_keys.count('{STATE}'), 1)
    check_header_count('{CNTYFILENAME}', county_sheet_keys.count('{CNTYFILENAME}'), 1)

    # check for one occurrence for keys for those we are checking max length
    for field in LENGTH_CHECK_FIELDS:
        check_header_count(f"{{{field}}}", county_sheet_keys.count(f"{{{field}}}"), 1)

    # check for one occurrence of field if we are checking length for a subset we are writing
    if LENGTH_CHECK_FIELDS_SELECT:
        check_header_count(f"{{{LENGTH_CHECK_FIELDS_SELECT}}}",
                           county_sheet_keys.count(f"{{{LENGTH_CHECK_FIELDS_SELECT}}}"), 1)

    # keys in file heading also being replaced in docxs
    county_sheet_required_keys = [c.upper() for c in county_sheet_keys if c.upper() in required_keys]

    # get key values and number of occurrences in list of keys
    key_counts = Counter(county_sheet_required_keys)

    bad_key_counts = [(k, key_counts[k]) for k in key_counts if key_counts[k] > 1]  # flag >1
    if bad_key_counts:
        fld_string = '\n'.join([f"{tup[0]}: {str(tup[1])}" for tup in bad_key_counts])
        msg = f"The following column keys are in the sheet more than once:\n\n{fld_string}"
        logger.error(msg)
        exit_yes(msg, "Duplicate Key Columns in Sheet")

    try:
        county_sheet = pd.read_excel(xlsx_file_name, sheet_name='Counties', header=0, skiprows=1).astype(str)
    except BaseException as e:
        msg = "Install openpyxl 3.0.10 not 3.1.2 because later bombs if xlsx has filters"
        logger.error(msg)
        logger.error(e)
        logger.error("Install openpyxl 3.0.10 not 3.1.2 because later bombs if xlsx has filters")
        exit_yes(msg)

    county_sheet.columns = [c.upper() for c in county_sheet.columns]  # convert all column names to uppercase
    # get list of required columns
    required_county_cols = [c.upper() for c in county_sheet.columns if c.upper() in required_keys]

    # only keep needed columns in df
    county_sheet = county_sheet[required_county_cols]
    # replace fields with stripped values then "' '" to account for and see multi spaces
    county_sheet = county_sheet.apply(lambda x: x.str.strip())

    # Create df of ony those counties being processed.  Used below.  Fields are cleaned up here.
    used_counties_df = county_sheet.loc[county_sheet['{USE}'].notnull()]

    # replace fields with string showing quotes ("' '") to make easier to see.
    # May not want to replace in file used for replacements so we could replace keys with spaces
    # in docxs
    used_counties_df.replace([""], ["' '"], inplace=True)
    list_of_used_county_names = used_counties_df['{CNTYFILENAME}'].to_list()
    string_of_used_county_names = ",".join(list_of_used_county_names)
    logger.info(f"Running on counties: {string_of_used_county_names}")
    logger.info("")
    exit_yes_no(f"Will process {used_counties_df.shape[0]} of {county_sheet.shape[0]} total counties:\n\n"
                f"{string_of_used_county_names}\n\nContinue?")

    # check if any fields to be substituted have nan values
    logger.debug("checking nans")
    nan_in_df = used_counties_df[used_counties_df.isnull().any(axis=1)]
    if not nan_in_df.empty:
        print(f"\nUSED COUNTIES CONTAINING 'nan' DATA")
        with pd.option_context('display.max_rows', None, 'display.max_columns', None):
            print(nan_in_df)
        pyautobek.alert("Some info to be merged contains 'nan'.\n\nSee log.", "ERROR: Used data contains 'nan'")

    # check if any fields to be substituted have blank values
    logger.debug("checking blanks")
    blank_in_df = used_counties_df[
        used_counties_df.eq("' '").any(axis=1)]  # all fields replaced with stripped val then ' '  above to see

    if not blank_in_df.empty:
        print(f"\nUSED COUNTIES CONTAINING DATA WITH SPACES")
        with pd.option_context('display.max_rows', None, 'display.max_columns', None):
            print(blank_in_df[list(keys_in_doc)])
        pyautobek.alert("Some info to be merged contains spaces.\n\nSee log.", "ERROR: Used data contains spaces")

    exit_yes_no("Close MS Word if open and continue?\n\nContinue?",
                'Continue or Exit?', )

    subprocess.call(['osascript', '-e', 'tell application "Word" to quit'])

    count = 0
    for _, county_row in used_counties_df.iterrows():  # uses df of only rows with used != nan from above
        state_fn = f"{county_row['{STATE}']}-{county_row['{CNTYFILENAME}']}"
        logger.debug("in used_counties_df.iterrows")

        for docx_file in docx_files:
            logger.debug("creating Document using Document(docx_file)")
            document_of_docx_file = Document(docx_file)
            script = deepcopy(document_of_docx_file)
            logger.debug("ready to call make_replacements")
            make_replacements(script, county_row)

            output_dir = path_root / 'Output'
            output_docxs_dir = path_root / 'Output' / 'Docxs'
            output_pdfs_dir = path_root / 'Output' / 'Pdfs'

            if not output_dir.exists():
                output_dir.mkdir()
            if not output_docxs_dir.exists():
                output_docxs_dir.mkdir()
            if not output_pdfs_dir.exists():
                output_pdfs_dir.mkdir()

            file_with_state = f"{state_fn} {Path(docx_file).name}"  # use filename var to match VL

            if SINGLE_COUNTY_OUTPUT_DIR:
                # script_filepath = output_dir / file_with_state
                script_filepath = output_docxs_dir / file_with_state
            else:
                state_county_dir = output_dir / state_fn  # added BEK 4/23/22 to use filename var
                if not state_county_dir.exists():
                    state_county_dir.mkdir()
                script_filepath = state_county_dir / file_with_state

            script.save(script_filepath)

        if True:  # convert docxs to pdfs
            if not SINGLE_COUNTY_OUTPUT_DIR:
                logger.debug("ready to sleep/pause")
                # time.sleep(2)  # No longer - Dropbox needs time to update or Word bombs
                logger.debug(f"calling convert on '{file_with_state}'")
                convert(state_county_dir)  # this converts all docx files in the folder to pdf using docs2pdf function
                logger.debug("after convert")

        count += 1
        msg = f"Finished filling Word templates for {state_fn}. Completed {count} of {used_counties_df.shape[0]} " \
              f"chosen counties."
        logger.info(msg)

    if True and SINGLE_COUNTY_OUTPUT_DIR:  # convert docxs to pdfs
        logger.debug("ready to sleep/pause")
        # time.sleep(2)  # No longer - Dropbox needs time to update or Word bombs
        logger.debug(f"calling convert on all docxs in Output directory")
        convert(output_docxs_dir, output_pdfs_dir)  # this converts all docx files in the folder to pdf using docs2pdf function
        logger.debug("after convert")

    msg = f"Completed scripts and Avery sheets for {count} counties total."
    logger.info(msg)
    pyautobek.alert(msg, "Alert")



def main():
    """ pick function from menu """
    import sys
    if "--update" in sys.argv:
        subprocess.run(["uvx", "--reinstall", "--from",
            "git+https://github.com/kramsman/uv_levitool.git", "levitool"])
        return

    choice = pyautobek.confirm("What do you want to do?", title='Select',
                      buttons=['LeviTool', 'Max Row Len', 'Check Urls', 'Check Short Urls', 'Exit'])
    choice = choice.lower()

    if choice == 'levitool':
        main_levi()
    elif choice == 'max row len':
        max_label_lengths(initial_attachment_dir=INITIAL_ATTACHMENT_DIR)
    elif choice == 'check urls':
        check_boe_urls(key_containing_url=None, key_for_used=None, open_browser=True)
    elif choice == 'check short urls':
        check_short_boe_urls(key_containing_short_url=None, key_containing_url=None, key_for_used=None,
                             open_browser=None)

    else:
        exit()


if __name__ == '__main__':

    main()
