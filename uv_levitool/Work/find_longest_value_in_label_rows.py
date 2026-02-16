""" Check the lengths of row text in a label after Levitool substitutions to make sure it will fit.
Should be max of 34.
"""

# FIXME check if {priority} changes the number of rows used in count

from copy import deepcopy
from pathlib import Path

from docx import Document  # package in Conda is python-docx, not simply docx
from uvbekutils import pyautobek
from uvbekutils import safe_str
from uvbekutils import scroll_box
from uvbekutils import list_pick


def get_label_text(label_docx):
    """ get text from upper leftmost cell of label and return a text string with lines separated by '\n'"""

    # string or path?
    document_of_docx_file = Document(label_docx)
    d = deepcopy(document_of_docx_file)

    # for line in d.tables[0].rows[0].cells[0].paragraphs:
    lines = [line.text for line in d.tables[0].rows[0].cells[0].paragraphs]
    label_text = '\n'.join(lines)
    # line = d.tables[0].rows[0].cells[0].paragraphs[0]
    print(f"{label_text=}")
    return label_text


def max_label_lengths(*, used_field: str = None, label_docx=None, initial_attachment_dir=None) -> None:
    """
    For checking length of lines in labels to be 'mail merged' by the Levitool. Length should be 34 or under.
    Given a block of text lines separated by return chars, this will make substitutions of df fields listed in
    field_list and print the longest substituted line and its length.

    Args:
        initial_attachment_dir ():
        label_docx (): template label used for substitutions.  Text is taken from upper left table cell.
        # label_text (): a block of text pasted from a label docx with lines separated by return chars
        # label_text_file (): a text file containing a block of text pasted from a label docx with lines separated by
        return chars.  ASSUMED TO BE NAMED label.txt IN SAME DIR AS BOE xlsx
        used_field (): field which represents used records in df, usually '{PRIORITY}'
    """

    import pandas as pd
    from pathlib import Path
    from loguru import logger
    import re
    import glob
    from uvbekutils import exit_yes, select_file

    from uv_levitool.Work.read_boe_xls import read_boe_xls

    def print_max_line_info(df: pd.DataFrame, line_list_field: str) -> str:
        """
        Calculates the longest line from a df cell containing a list of the substituted text, its length and prints
        them.

        Args:
            df (): df with field info to be substituted in label text
            line_list_field (): the field in df which contains a list of the lines with text substituted

        Returns:
            str: formatted string with max line info for each line
        """
        if df.empty:
            pyautobek.alert(f"No counties are being selected based on the field '{used_field}'.\n\n"
                            f"Checking for these copunties will be skipped.\n"
                            f"If desired, rerun choosing another field (like {{priority}}'" ,
                        "No Counties Being Selected")
            result = 'None'
        else:
            # get number of list elements (# of lines) in the list field, so we know how many lines need to be checked
            number_of_lines = len(df[line_list_field].iloc[0])
            result_lines = []
            for line_number in range(number_of_lines):
                line_data = [line[line_number] for line in df[line_list_field]]
                max_line = max(line_data, key=len)
                max_line_length = len(max_line)
                # print(f"line_data: {line_data}")
                # line_info = f"line {line_number + 1} max_line: '{max_line}', len:{max_line_length}"
                line_info = f"line {line_number + 1},  max: {max_line_length}   '{max_line}'"
                print(line_info)
                result_lines.append(line_info)
                result = '\n'.join(result_lines)

        return result

    if label_docx is None:
        label_docx = select_file("PICK LABEL DOCX",
                                 start_dir=initial_attachment_dir,
                                 files_like='*.docx',
                                 mode='file',
                                 title2="Pick label docx template that will have BOE info merged into it "
                                        "('LABEL 30 per page...')",
                                 )
    logger.info(f"Label docx being used for input: '{label_docx}")
    label_docx = Path(label_docx).expanduser()

    label_text = get_label_text(label_docx)
    keys = find_keys_in_text(label_text)
    keys = [key.lower() for key in keys]

    # FIXME should tis be live?
    # if boe_xls is None:
    #     boe_xls = get_file_name("Pick BOE file (xlsx) with urls to check.",
    #                              initial_dir=label_docx.parent,
    #                              title2="Pick BOE file (xlsx)")
    # logger.info(f"BOE xls being checked is: '{boe_xls}")
    # list of all xlsx files not starting with ~ (temporary files)
    xlsx_files = glob.glob(str(label_docx.parent / '[!~]*.xlsx'))

    if len(xlsx_files) != 1:  # must have 1, not more
        msg = f"Can not select BOE xlsx file from label directory because there are more than one;" \
              f"there are {len(xlsx_files)}.  EXITING."
        exit_yes(msg, "MORE THAN ONE XLSX")
    # xlsx_file_name = xlsx_files[0]  # filename of the first (and only) xlsx
    boe_xls = Path(xlsx_files[0]).expanduser()

    county_sheet_keys, df = read_boe_xls(boe_xls)

    # Lowercase all column names to match lowercase keys
    df.columns = [col.lower() for col in df.columns]


    if used_field is None:
        used_field = list_pick(lst=df.columns,
                               title="PICK KEY FIELD TO ID 'USED' COUNTIES",
                               msg="All counties with this field not empty will be checked as a group ("
                                   "'{priority}' will be used if none chosen).",
                               select_mode='single',
                               pre_select=False,
                               allow_none=True,
                               )[0]
        if used_field == '':
            used_field = '{priority}'

    logger.info(f"Label docx being used for input: '{label_docx}")


    # set object type so field can accept an object - a list
    df['lines'] = ''
    df['lines'] = df['lines'].astype('object')

    for index, row in df.iterrows():
        tmp_label_text = label_text  # so we start fresh with the label info containing the {} keys

        for fld in keys:
            # replace occurrences of the key with the value in the df (ex '{county}' gets replaced with 'Dade')
            # tmp_label_text = tmp_label_text.replace(fld, df.at[index, fld])
            # TODO what ot do if nan?  takes up only one space but might be much larger when filled in later.
            tmp_label_text = re.sub(fld, safe_str(df.at[index, fld]), tmp_label_text, flags=re.IGNORECASE)

        line_list = safe_str(tmp_label_text).split('\n')  # create list of lines from the text
        line_list = [line.strip('\t ') for line in line_list if len(line.strip('\t ')) > 0]  # trim and remove blanks

        # fill a field in the df with the list of line text
        df.at[index, 'lines'] = line_list

    # Collect max line info for alert display
    max_line_results = []
    total_rows = len(df)

    # Below is True if 'nan' found in any cell in 'keys' columns
    nan_found_in_all_rows = (df[keys] == 'nan').any(axis="columns").any(axis="rows")
    # Below creates a boolean series/mask: (df[keys] == 'nan').any(axis="columns")
    # df[keys] selects only columns in list 'keys'
    # below puts it together: creates df with 'keys' columns and rows containing 'nan' in any column:
    #   df[keys][(df[keys] == 'nan').any( axis="columns")]
    df_nan = df[keys][(df[keys] == 'nan').any(axis="columns")]

    if nan_found_in_all_rows:
        print("\n20 rows containing 'nan' somewhere in a ALL row 'key' columns")
        with pd.option_context('display.max_rows', 20, 'display.max_columns', None, ):
            print(df_nan)
        pyautobek.alert(f"'Nan' (empty) found in {len(df_nan)} rows in at least one key of '{keys}' to be substituted "
                        f"in ALL rows so analysis of all rows skipped.\n\nSee log for print of 20 rows.",
                        "'nan' found somewhere in ALL row keys")

        # there are some keys with nan so can't run on all.  Instead, find those rows with all keys non-nan.
        df_non_nan = df[(df[keys].notnull()).all(axis="columns")]  # note all columns taken
        if len(df_non_nan) != 0:
            print("ALL ROWS WITH NO NAN VALUES")
            print(df_non_nan[keys])
            print()
            print("MAX LINES IN SUBSTITUTED SCRIPT INFO FOR ALL ROWS WITH NO NAN VALUES")
            result = print_max_line_info(df_non_nan, 'lines')
            max_line_results.append(("ALL ROWS WITH NO NAN VALUES", result, len(df_non_nan)))
        print()
        a = 1
    else:
        print("MAX LINES IN SUBSTITUTED SCRIPT INFO FOR ALL ROWS")
        result = print_max_line_info(df, 'lines')
        max_line_results.append(("ALL ROWS", result, len(df)))
        print()

    # see explanation of expression above
    # Check for bad row where NAN is found somewhere in a key field.
    nan_found_in_select_rows = (df[keys].loc[df[used_field].str.strip().notnull()] == 'nan') \
        .any(axis="columns").any(axis="rows")
    df_nan = df[keys][(df[used_field].str.strip().notnull()) & (df[keys] == 'nan').any(axis="columns")]
    if nan_found_in_select_rows:
        print(f"\nThe following rows with non-blank '{used_field}' contain 'nan' somewhere in ALL 'key' columns")
        with pd.option_context('display.max_rows', None, 'display.max_columns', None, ):
            print(df_nan)

        print(f"\nSTATS NOT PRODUCED - USED ROWS HAVE BAD 'nan' DATA")

        pyautobek.alert(
            f"'Nan' (blank) found in {len(df_nan)} rows at least one key with non-blank '{used_field}' to be substituted so analysis "
            f"of all rows skipped.\n\n"
            "See log for list for rows.",
            "'nan' found in USED row keys")
    else:
        print(f"MAX LINES IN SUBSTITUTED SCRIPT INFO FOR USED SUB ROWS ('{used_field}' not blank)")
        df_used_counties = df.loc[df[used_field].str.strip().notnull()]
        result = print_max_line_info(df_used_counties, 'lines')
        max_line_results.append((f"'USED' COUNTIES ('{used_field}' not blank)", result, len(df_used_counties)))
        print()

    print("RAW LABEL DATA:")
    print(label_text)

    # Build and display alert with max line lengths and label text
    alert_lines = []
    alert_lines.append("--- LABEL TEMPLATE TEXT ---")
    alert_lines.append(label_text)
    alert_lines.append("")
    for section_name, section_result, row_count in max_line_results:
        alert_lines.append(f"--- {section_name} ---")
        alert_lines.append(section_result)
        alert_lines.append(f"({row_count} rows of {total_rows})")
        alert_lines.append("")

    alert_lines.append(f"'Used' counties based on field {used_field}")
    alert_lines.append("")
    alert_lines.append(f"Label File: {label_docx}")
    alert_lines.append("")
    alert_lines.append(f"BOE Sheet: {boe_xls}")
    alert_lines.append("")

    # Add font size reference tables
    alert_lines.append("")
    alert_lines.append("")
    alert_lines.append("Estimated character limits")
    alert_lines.append("")
    alert_lines.append("30 per page, Avery 5161 Label")
    alert_lines.append("Font Size | Approximate Characters per Line")
    alert_lines.append("----------|--------------------------------")
    alert_lines.append("8 pt        52-55 characters")
    alert_lines.append("9 pt        46-49 characters")
    alert_lines.append("10 pt       42-45 characters")
    alert_lines.append("11 pt       38-41 characters")
    alert_lines.append("12 pt       35-38 characters")
    alert_lines.append("13 pt       32-35 characters")
    alert_lines.append("14 pt       30-33 characters")
    alert_lines.append("")
    alert_lines.append("20 per page, Avery 5161 Label")
    alert_lines.append("Font Size | Approximate Characters per Line")
    alert_lines.append("----------|--------------------------------")
    alert_lines.append("8 pt        80-85 characters")
    alert_lines.append("9 pt        71-76 characters")
    alert_lines.append("10 pt       64-68 characters")
    alert_lines.append("11 pt       58-62 characters")
    alert_lines.append("12 pt       53-57 characters")
    alert_lines.append("13 pt       49-52 characters")
    alert_lines.append("14 pt       46-49 characters")

    alert_message = '\n'.join(alert_lines)
    # pyautobek.alert(alert_message, "MAX LABEL LINE LENGTHS")
    scroll_box(alert_message, title="MAX LABEL LINE LENGTHS", wrap_lines=True )


def find_keys_in_text(label_text):
    """ Finds keys, fields in document to be replaced, in flat text"""

    import re
    keys = set(re.findall(r'{[a-zA-Z0-9]+}', safe_str(label_text)))
    return keys


if __name__ == '__main__':
    INITIAL_ATTACHMENT_DIR = Path("~/Dropbox/Postcard Files/Attachments/Campaigns").expanduser()
    # LENGTH_CHECK_FIELDS_SELECT = '{priority}'
    LENGTH_CHECK_FIELDS_SELECT = '{use}'
    max_label_lengths(used_field=None, initial_attachment_dir=INITIAL_ATTACHMENT_DIR, )

    a = 1
    # Label text is copied here.  Lines are trimmed.
    #     LABEL_TEXT = """
    # {cntytoprint} Registrar
    # Phone Num: {specialphone}
    # or {specialurl}
    # Early Voting Now thru Feb 18
    #     """

    # 30 per page, Avery 5161 Label
    # Font Size | Approximate Characters per Line
    # ----------|--------------------------------
    # 8 pt      | 52-55 characters
    # 9 pt      | 46-49 characters
    # 10 pt     | 42-45 characters
    # 11 pt     | 38-41 characters
    # 12 pt     | 35-38 characters
    # 13 pt     | 32-35 characters
    # 14 pt     | 30-33 characters

    # 20 per page, Avery 5161 Label
    # Font Size | Approximate Characters per Line
    # ----------|--------------------------------
    # 8 pt      | 80-85 characters
    # 9 pt      | 71-76 characters
    # 10 pt     | 64-68 characters
    # 11 pt     | 58-62 characters
    # 12 pt     | 53-57 characters
    # 13 pt     | 49-52 characters
    # 14 pt     | 46-49 characters

    # label_text_file = Path("~/Dropbox/Postcard Files/Attachments/Campaigns/TEST NAN3/Input/label.txt").expanduser()
    # with open(label_text_file, "r") as f:
    #     LABEL_TEXT = f.read()

    # LENGTH_CHECK_FIELDS_SELECT = '{priority}'
    # LENGTH_CHECK_FIELDS = [['county', 'phone'], 'county', 'phone']
    # LENGTH_CHECK_FIELDS = ['{cntytoprint}', '{specialphone}', '{specialurl}']
    # test = ['+'.join(x) for x in LENGTH_CHECK_FIELDS]
    # test = ['county', 'phone']

    # data = [['Small', '555-555-5555', 'x'], ['Medium', '555-555-5552 x2', 'x'], ['VeryVeryLonnnng', '555-555-5553 ext 3', '']]
    # df = pd.DataFrame(data, columns=['{county}', '{phone}', '{PRIORITY}'])

    # max_label_lengths(initial_attachment_dir=INITIAL_ATTACHMENT_DIR,)
    # label_docx="~/Dropbox/Postcard Files/Attachments/Campaigns/TEST NAN3/Input/TEST Special GOTV "
    #            "LABELS-30per page 1-3-2023.docx")
    # a = 1

    # boe_xls = "~/Dropbox/Postcard Files/Attachments/Campaigns/TEST NAN3/Input/BOE Info VA (1).xlsx",

    # df_to_check_length = pd.DataFrame
    #
    # fieldnames = []
    # for item in LENGTH_CHECK_FIELDS:
    #     if isinstance(item, list):
    #         fieldname = '+'.join(item)
    #         df_to_check_length[fieldname] = df[item].astype(str).apply(lambda x: ''.join(x), axis=1)
    #     elif isinstance(item, str):
    #         df_to_check_length[item] = df[item]
    #     else:
    #         print('ERROR - field not on dataframe')
    #     fieldnames.append(fieldname)
    #
    # df_to_check_length[LENGTH_CHECK_FIELDS_SELECT] = df_to_check_length[LENGTH_CHECK_FIELDS_SELECT]
    #
    #
    # for field in LENGTH_CHECK_FIELDS:
    #     # max_field = max(df[f"{{{field}}}"], key=len)
    #     max_field = max(df[f"{field}"], key=len)
    #     print(f"Max of '{field}' column: '{max_field}', length is {len(max_field)}")
    #
    # a=1

    # def fill_lines(df, field_list, used_field):
    #
    #     # set object type so field can accept an object - a list
    #     df['lines'] = ''
    #     df['lines'] = df['lines'].astype('object')
    #
    #     for fld in field_list:
    #         locals()[fld] = fld
    #         # newfield = df[fld]
    #
    #     for index, row in df.iterrows():
    #         # county = row['county']
    #         # phone = row['phone']
    #         # lines.append(f"{county} is the county name.")
    #
    #         df.at[index, 'lines'] = \
    #             [
    #                 f"{county} is the county name.",
    #                 f"{phone} is the phone.",
    #                 f"Together we have {county} and {phone}.",
    #                 f"Again we have {county} and {phone}."
    #             ]
    #
    #     entries = len(df.at[0, 'lines'])
    #
    #     for entry in range(entries):
    #         line_data = [x[entry] for x in df['lines'] ]
    #         max_line = max(line_data, key=len)
    #         max_length = len(max_line)
    #         print(f"line_data: {line_data}")
    #         print(f"max_line: '{max_line}', len:{max_length}\n")
    #
    #
    #     # for entry in range(entries):
    #     #     # max_field = max(df[f"{{{field}}}"], key=len)
    #     #     max_field = max(df['line'][entry], key=len)
    #     #     print(f"Max of '{field}' column: '{max_field}', length is {len(max_field)}")
    #
    a = 1
