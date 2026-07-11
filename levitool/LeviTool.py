"""
LeviTool.py creates the attachments need for postcard writing namely scripts and labels.  It does this by performing
a "mail merge" on input files, replacing fields in {} with those from an Excel input file.
All the input files must be in a directory named 'Input', and must contain one xlsx with the info to be merged and
any number of docx files which will have the dat merged into them.
A pdf of each 'docx' is created after the merge.
A separate subdirectory under a min directory named "Output" is created for each used row in the input xlsx file.

To run this in Pycharm, use a configuations set to 'module' with a name like "levitool-root.LeviTool"
and a working dir like "/Users/Denise/Library/CloudStorage/Dropbox/non ROV python progs/levitool-root"
"""

# # After creating your project with uv init to create requirements.txt
# uv pip compile pyproject.toml -o requirements.txt

# NEED TO USE BACK VERSION OF OPENPYXL for bug caused in 3.1: county_sheet_df = pd.ExcelFile( script_file_name).parse(
# use openpyxl 3.0.10 not 3.1.2 because later bombs if xlsx has filters  https://stackoverflow.com/questions/75382340/python-pandas-read-excel-error-value-must-be-either-numerical-or-a-string-conta

# TODO currently blanks replaced with "' '" in df used_counties_df.  OK or should be blank? Or '.'?
# TODO Check max length of fields not working. But it's in another version - levitool in non ROV Python?

import sys
import os
sys.path.append(os.path.expanduser("~/Dropbox/Postcard Files/"))

import subprocess
import glob
import re
import shutil
import tempfile
from collections import Counter
from copy import deepcopy
from pathlib import Path

import pandas as pd
from uvbekutils import exit_yes, exit_yes_no, setup_loguru, select_file
from uvbekutils import pyautobek
from .check_boe_urls import check_boe_urls, check_short_boe_urls
from docx import Document  # package in Conda is python-docx, not simply docx
from loguru import logger
from .find_longest_value_in_label_rows import max_label_lengths
from .constants import (INITIAL_ATTACHMENT_DIR, PROGRAM_REQUIRED_KEYS, TEST_INPUT_DIR,
                        TEST_SKIP_PROMPTS, DOC_TO_PDF_CONVERTER)

pd.options.mode.copy_on_write = True  # fix chain assignment forced in Pandas 3.0

# NOTE: setup_loguru is called once inside main(), not at module load. Running it at import
# adds a second log sink (hence duplicate log lines) when this module is executed twice —
# e.g. `python -m levitool.LeviTool` imports it as both 'levitool.LeviTool' and '__main__'.


def count_brackets(s: str) -> int:
    """Counts the net difference between left and right curly brackets in a string.

    Args:
        s (str): The string to examine for curly brackets.

    Returns:
        int: Net bracket count (left minus right), or -10000 if the absolute
            difference exceeds 1, indicating a likely template formatting error.
    """

    bracket_count = s.count('{') - s.count('}')
    if bracket_count < -1 or bracket_count > 1:
        bracket_count = -10000

    return bracket_count


def find_keys(document) -> set:
    """Finds all {KEY} replacement tokens in a docx document.

    Searches section headers, body paragraphs, and all table cells (three levels deep).

    Args:
        document: A python-docx Document object to scan for replacement keys.

    Returns:
        set: Set of uppercase key strings found (e.g. '{COUNTY}', '{STATE}').
    """
    # Keys are fields in document to be replaced
    keys = set()

    d = deepcopy(document)

    for section in d.sections:
        for paragraph in section.header.paragraphs:
            new_keysx = re.findall(r'{[a-zA-Z0-9]+}', paragraph.text)
            new_keys = [x.upper() for x in new_keysx]
            keys = keys.union({k.upper() for k in new_keys})

    for paragraph in d.paragraphs:
        new_keysx = re.findall(r'{[a-zA-Z0-9]+}', paragraph.text)
        new_keys = [x.upper() for x in new_keysx]
        keys = keys.union({k.upper() for k in new_keys})

    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    new_keysx = re.findall(r'{[a-zA-Z0-9]+}', paragraph.text)
                    new_keys = [x.upper() for x in new_keysx]
                    keys = keys.union({k.upper() for k in new_keys})

    return keys


def replace_tags(paragraph, county_row: pd.Series) -> None:
    """Replaces all {KEY} tokens in a paragraph's runs with values from the county row.

    Handles keys that span multiple runs due to docx formatting by reassembling
    partial runs into a single string before performing substitution.

    Args:
        paragraph: A python-docx Paragraph object whose runs will be modified in place.
        county_row (pd.Series): A row from the county DataFrame mapping keys to
            replacement values.

    Raises:
        ValueError: If bracket nesting in the paragraph runs is inconsistent,
            indicating a template formatting error near the offending run text.
    """

    brackets = 0
    partial_runs = []
    for run in paragraph.runs:
        partial_runs.append(run)
        brackets += count_brackets(run.text)
        if brackets == 0:
            text = ''.join([r.text for r in partial_runs])
            for key, val in county_row.items():
                text = re.sub(key, val, text, flags=re.IGNORECASE)  # use re to replace and ignore case
            partial_runs[0].text = text
            for r in partial_runs[1:]:
                r.text = ''

            brackets = 0
            partial_runs = []
        elif brackets != 1:
            msg = f'Template error near "{run.text}" in script template. Make sure sequence to be replaced has consistent formatting.'
            logger.error(msg)
            raise ValueError(msg)


def make_replacements(document, county_row: pd.Series) -> None:
    """Applies tag replacement across all paragraphs and table cells in a document.

    Args:
        document: The python-docx Document to modify in place.
        county_row (pd.Series): A row from the county DataFrame mapping keys to
            replacement values.
    """

    for paragraph in document.paragraphs:
        replace_tags(paragraph, county_row)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_tags(paragraph, county_row)

    return


def check_header_count(hdr_str: str, header_count: int, num_allowed: int) -> None:
    """Validates that a header token appears exactly the expected number of times; exits on failure.

    Args:
        hdr_str (str): The header token being checked (e.g. '{USE}').
        header_count (int): The actual number of times the header was found.
        num_allowed (int): The expected number of occurrences.
    """

    if header_count != num_allowed:
        msg = f"'{hdr_str}' does not appear {num_allowed} time(s), instead {header_count}\n\nEXITING."
        logger.error(msg)
        exit_yes(msg, "ERROR")


def select_input_directory(initial_dir: Path) -> tuple:
    """Prompts for an Input directory via file picker, or uses a hardcoded path if TEST_INPUT_DIR is True.

    Validates that the selected directory name ends with 'Input'.

    Args:
        initial_dir (Path): Starting directory for the file picker dialog.

    Returns:
        tuple: A (path_input, path_root) pair where path_input is the selected Input
            directory and path_root is its parent.
    """

    if not TEST_INPUT_DIR:
        path_input = select_file(
            title="Pick 'Input' Directory",
            start_dir=initial_dir,
            files_like="Input",
            choices=["Select", "Cancel"],
            mode="dir",
            title2="Select the directory which contains all your template files (docs and the xlsx). Must be named 'Input'"
        )
        print(f"Selected: {path_input}")
        path_input = Path(path_input)
    else:
        path_input = Path("~/Dropbox/Postcard Files/Attachments/Campaigns/Test box and image/Input").expanduser()

    path_root = path_input.parents[0]

    if path_input.parts[-1].upper() != "INPUT":
        msg = f"Last part of chosen directory:\n\n {path_input}\n\nDoes not equal 'INPUT'.\n\nEXITING."
        logger.error(msg)
        exit_yes(msg, "ERROR")

    return path_input, path_root


def load_template_files(path_input: Path) -> tuple:
    """Finds and validates the single xlsx and one or more docx templates in the Input directory.

    Exits with an error message if there is not exactly one xlsx or no docx files.

    Args:
        path_input (Path): The Input directory to search for template files.

    Returns:
        tuple: A (xlsx_file, docx_files) pair where xlsx_file is the path string to
            the xlsx and docx_files is a list of path strings to the docx templates.
    """

    xlsx_files = glob.glob(str(path_input / '[!~]*.xlsx'))
    if len(xlsx_files) != 1:
        msg = (f"Must have only 1 xlsx file in the Input directory; there are {len(xlsx_files)}.  \n\n"
               f"{path_input}\n\n"
               f"EXITING.")
        logger.error(msg)
        exit_yes(msg, "ERROR")
    xlsx_file = xlsx_files[0]

    docx_files = glob.glob(str(path_input / '[!~]*.docx'))
    if not docx_files:
        msg = f"No docx input file found in directory: '{path_input}' EXITING."
        logger.error(msg)
        exit_yes(msg, "ERROR")

    logger.info(f"Will replace tags in {len(docx_files)} documents per county.\n")

    return xlsx_file, docx_files


def extract_required_keys(docx_files: list) -> set:
    """Collects all {KEY} replacement tokens from all docx templates plus program-required keys.

    Adds mandatory keys ({STATE}, {CNTYFILENAME}, {USE}) to the set gathered from the templates.

    Args:
        docx_files (list): List of file path strings to docx template files.

    Returns:
        set: Set of uppercase key strings required across all templates and program logic.
    """

    required_keys = set()
    for docx_file in docx_files:
        document_of_docx_file = Document(docx_file)
        keys_in_doc = find_keys(document_of_docx_file)
        required_keys = required_keys.union(keys_in_doc)

    # Add keys required by the program but not necessarily in the docx templates
    required_keys.update(PROGRAM_REQUIRED_KEYS)

    return required_keys


def load_and_validate_excel(xlsx_file: str, required_keys: set) -> pd.DataFrame:
    """Reads the Counties sheet, validates headers and duplicate keys, and returns a filtered DataFrame.

    Checks that mandatory tokens ({USE}, {STATE}, {CNTYFILENAME}) appear exactly once,
    then reads the full sheet and filters to only the required columns.

    Args:
        xlsx_file (str): Path to the xlsx file containing the Counties sheet.
        required_keys (set): Set of uppercase key strings that must exist in the sheet.

    Returns:
        pd.DataFrame: DataFrame containing only the required columns, with all values
            stripped of whitespace and read as strings.
    """

    logger.info(f"Required Levitool keys: {required_keys}")

    # Read just the header row to get keys
    county_sheet_df = pd.read_excel(xlsx_file, sheet_name='Counties', header=None, skiprows=1,
                                    nrows=1).astype(str)
    county_sheet_keys = county_sheet_df.loc[0, :].values.tolist()
    logger.info(f"Keys contained in county tab of BOE file: {county_sheet_keys}")
    for key in county_sheet_keys:
        logger.debug(f"county_sheet_keys field: value={key!r}, type={type(key)}")

    county_sheet_keys = [x.upper() for x in county_sheet_keys if not pd.isna(x)]

    # Check that required tokens appear exactly once
    check_header_count('{USE}', county_sheet_keys.count('{USE}'), 1)
    check_header_count('{STATE}', county_sheet_keys.count('{STATE}'), 1)
    check_header_count('{CNTYFILENAME}', county_sheet_keys.count('{CNTYFILENAME}'), 1)

    # Check for duplicate required keys
    county_sheet_required_keys = [c.upper() for c in county_sheet_keys if c.upper() in required_keys]
    key_counts = Counter(county_sheet_required_keys)
    bad_key_counts = [(k, key_counts[k]) for k in key_counts if key_counts[k] > 1]
    if bad_key_counts:
        fld_string = '\n'.join([f"{tup[0]}: {str(tup[1])}" for tup in bad_key_counts])
        msg = f"The following column keys are in the sheet more than once:\n\n{fld_string}"
        logger.error(msg)
        exit_yes(msg, "Duplicate Key Columns in Sheet")

    # Read full sheet
    try:
        county_sheet = pd.read_excel(xlsx_file, sheet_name='Counties', header=0, skiprows=1).astype(str)
    except BaseException as e:
        msg = "Install openpyxl 3.0.10 not 3.1.2 because later bombs if xlsx has filters"
        logger.error(msg)
        logger.error(e)
        exit_yes(msg)
        raise  # unreachable; satisfies linter that county_sheet is always assigned

    county_sheet.columns = [c.upper() for c in county_sheet.columns]
    required_county_cols = [c.upper() for c in county_sheet.columns if c.upper() in required_keys]
    county_sheet = county_sheet[required_county_cols]
    county_sheet = county_sheet.apply(lambda x: x.str.strip())

    return county_sheet


def filter_and_check_data(county_sheet: pd.DataFrame) -> pd.DataFrame:
    """Filters the county sheet to used rows and checks for nan or blank values.

    Prompts the user to confirm the county selection and to close MS Word, unless
    TEST_SKIP_PROMPTS is True. Displays popup alerts if nan or blank values are found
    in used rows.

    Args:
        county_sheet (pd.DataFrame): Full county DataFrame with all columns including {USE}.

    Returns:
        pd.DataFrame: Filtered DataFrame containing only rows where {USE} is not null or empty.
    """

    used_counties_df = county_sheet.loc[county_sheet['{USE}'].notnull() & (county_sheet['{USE}'] != "")]
    used_counties_df.replace([""], ["' '"], inplace=True)

    list_of_used_county_names = used_counties_df['{CNTYFILENAME}'].to_list()
    string_of_used_county_names = ",".join(list_of_used_county_names)
    logger.info(f"Running on counties: {string_of_used_county_names}")
    logger.info("")

    if not TEST_SKIP_PROMPTS:
        exit_yes_no(f"Will process {used_counties_df.shape[0]} of {county_sheet.shape[0]} total counties:\n\n"
                    f"{string_of_used_county_names}\n\nContinue?")

    # Check for nan values
    logger.debug("checking nans")
    nan_in_df = used_counties_df[used_counties_df.isnull().any(axis=1)]
    if not nan_in_df.empty:
        print(f"\nUSED COUNTIES CONTAINING 'nan' DATA")
        with pd.option_context('display.max_rows', None, 'display.max_columns', None):
            print(nan_in_df)
        pyautobek.alert("Some info to be merged contains 'nan'.\n\nSee log.", "ERROR: Used data contains 'nan'")

    # Check for blank values
    logger.debug("checking blanks")
    blank_in_df = used_counties_df[used_counties_df.eq("' '").any(axis=1)]
    if not blank_in_df.empty:
        print(f"\nUSED COUNTIES CONTAINING DATA WITH SPACES")
        with pd.option_context('display.max_rows', None, 'display.max_columns', None):
            print(blank_in_df)
        pyautobek.alert("Some info to be merged contains spaces.\n\nSee log.", "ERROR: Used data contains spaces")

    if not TEST_SKIP_PROMPTS:
        exit_yes_no("This will close MS Word if open\n\nContinue?", 'Continue or Exit?')

    return used_counties_df


def setup_output_dirs(path_root: Path) -> tuple:
    """Creates Output, Output/Docxs, and Output/Pdfs directories under path_root if they do not exist.

    Args:
        path_root (Path): The root campaign directory (parent of the Input directory).

    Returns:
        tuple: A (output_docxs_dir, output_pdfs_dir) pair of Path objects for the
            docx and PDF output directories.
    """

    output_dir = path_root / 'Output'
    output_docxs_dir = path_root / 'Output' / 'Docxs'
    output_pdfs_dir = path_root / 'Output' / 'Pdfs'

    for d in [output_dir, output_docxs_dir, output_pdfs_dir]:
        if not d.exists():
            d.mkdir()

    return output_docxs_dir, output_pdfs_dir


def generate_documents(used_counties_df: pd.DataFrame, docx_files: list, output_docxs_dir: Path) -> int:
    """Quits MS Word, then fills all docx templates for each used county row and saves them.

    For each county row, makes a deep copy of each template, substitutes all {KEY}
    tokens with county data, and saves the result named '{STATE}-{CNTYFILENAME} {template_name}'.

    Args:
        used_counties_df (pd.DataFrame): DataFrame of counties to process (pre-filtered to {USE} rows).
        docx_files (list): List of file path strings to docx template files.
        output_docxs_dir (Path): Directory where the filled docx files will be saved.

    Returns:
        int: Number of counties processed.
    """

    subprocess.call(['osascript', '-e', 'tell application "Word" to quit'])

    count = 0
    for _, county_row in used_counties_df.iterrows():
        state_fn = f"{county_row['{STATE}']}-{county_row['{CNTYFILENAME}']}"
        logger.debug("in used_counties_df.iterrows")

        for docx_file in docx_files:
            logger.debug("creating Document using Document(docx_file)")
            document_of_docx_file = Document(docx_file)
            script = deepcopy(document_of_docx_file)
            logger.debug("ready to call make_replacements")
            make_replacements(script, county_row)

            file_with_state = f"{state_fn} {Path(docx_file).name}"
            script_filepath = output_docxs_dir / file_with_state
            script.save(script_filepath)

        count += 1
        msg = (f"Finished filling Word templates for {state_fn}. "
               f"Completed {count} of {used_counties_df.shape[0]} chosen counties.")
        logger.info(msg)

    return count


def convert_to_pdfs(output_docxs_dir: Path, output_pdfs_dir: Path) -> None:
    """Converts all docx files in the output docx directory to PDFs.

    Dispatches on the module-level ``DOC_TO_PDF_CONVERTER`` flag:
    ``"docx2pdf"`` uses Word's Save-As-PDF exporter (fast, but drops color-font
    glyphs); ``"wordexporter"`` drives Word's Print -> Save as PDF so COLR/CPAL
    color symbols are preserved; ``"libreoffice"`` converts headless via
    LibreOffice, which also preserves the color symbol with no template changes.

    Args:
        output_docxs_dir (Path): Directory containing the filled docx files to convert.
        output_pdfs_dir (Path): Directory where the resulting PDF files will be saved.

    Raises:
        ValueError: If ``DOC_TO_PDF_CONVERTER`` is not a recognized value.
    """

    logger.debug(f"converting docxs to pdfs via '{DOC_TO_PDF_CONVERTER}'")
    if DOC_TO_PDF_CONVERTER == "docx2pdf":
        from docx2pdf import convert  # lazy: only needed for this path
        convert(output_docxs_dir, output_pdfs_dir)
    elif DOC_TO_PDF_CONVERTER == "wordexporter":
        _convert_via_wordexporter(output_docxs_dir, output_pdfs_dir)
    elif DOC_TO_PDF_CONVERTER == "libreoffice":
        _convert_via_libreoffice(output_docxs_dir, output_pdfs_dir)
    else:
        raise ValueError(
            f"Unknown DOC_TO_PDF_CONVERTER {DOC_TO_PDF_CONVERTER!r}; "
            f"expected 'docx2pdf', 'wordexporter', or 'libreoffice'.")
    logger.debug("after convert")


def _convert_via_wordexporter(output_docxs_dir: Path, output_pdfs_dir: Path) -> None:
    """Converts each docx to PDF through Word's Print -> "Save as PDF" path.

    Runs the bundled ``word_print_to_pdf.jxa`` once per docx via ``osascript``.
    Unlike ``docx2pdf`` (Word's exporter), this routes through macOS
    Quartz/CoreText and therefore preserves COLR/CPAL color-font glyphs such as
    the CFCG color symbol. Word's default print filename equals the document
    name, so each file lands at ``<output_pdfs_dir>/<stem>.pdf``, matching the
    ``docx2pdf`` naming contract.

    Args:
        output_docxs_dir (Path): Directory containing the filled docx files to convert.
        output_pdfs_dir (Path): Directory where the resulting PDF files will be saved.

    Raises:
        RuntimeError: If the JXA helper reports an error or a PDF is not produced.
    """

    jxa = Path(__file__).parent / "word_print_to_pdf.jxa"
    docx_files = sorted(output_docxs_dir.glob("*.docx"))
    for i, docx_file in enumerate(docx_files, start=1):
        expected_pdf = output_pdfs_dir / (docx_file.stem + ".pdf")
        logger.debug(f"wordexporter: printing {docx_file.name} ({i}/{len(docx_files)})")
        result = subprocess.run(
            ["/usr/bin/osascript", "-l", "JavaScript", str(jxa),
             str(docx_file), str(output_pdfs_dir)],
            capture_output=True, text=True)
        message = (result.stdout or "").strip()
        if result.returncode != 0 or not message.startswith("OK"):
            raise RuntimeError(
                f"wordexporter failed for {docx_file.name}: "
                f"{message or result.stderr.strip()}")
        if not expected_pdf.exists():
            raise RuntimeError(
                f"wordexporter reported success but {expected_pdf.name} was not created")
    logger.info(f"wordexporter converted {len(docx_files)} docx files to PDF")


def _find_soffice() -> str:
    """Locates the LibreOffice ``soffice`` executable across platforms.

    Returns:
        str: Path to (or name of) the ``soffice`` executable.

    Raises:
        FileNotFoundError: If LibreOffice is not installed / not found.
    """

    candidates = [
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",   # macOS
        "soffice",                                                # Linux / on PATH
        "libreoffice",                                            # some Linux distros
        r"C:\Program Files\LibreOffice\program\soffice.exe",      # Windows
    ]
    for candidate in candidates:
        found = shutil.which(candidate) or (candidate if Path(candidate).exists() else None)
        if found:
            return found
    raise FileNotFoundError(
        "LibreOffice 'soffice' not found. Install LibreOffice to use "
        "DOC_TO_PDF_CONVERTER='libreoffice' (e.g. 'brew install --cask libreoffice').")


def _convert_via_libreoffice(output_docxs_dir: Path, output_pdfs_dir: Path) -> None:
    """Converts all docx files to PDF headlessly via LibreOffice.

    Runs a single ``soffice --headless --convert-to pdf`` over the whole
    directory. Unlike ``docx2pdf`` (Word's exporter), LibreOffice renders
    COLR/CPAL color-font glyphs, so the CFCG color symbol survives with no
    template changes. Note that LibreOffice lays out DOCX slightly differently
    from Word, so label wrapping/margins should be verified once. Output names
    match the ``docx2pdf`` contract (``<stem>.pdf`` in ``output_pdfs_dir``).

    Args:
        output_docxs_dir (Path): Directory containing the filled docx files to convert.
        output_pdfs_dir (Path): Directory where the resulting PDF files will be saved.

    Raises:
        FileNotFoundError: If LibreOffice is not installed.
        RuntimeError: If the conversion fails or a PDF is not produced.
    """

    soffice = _find_soffice()
    docx_files = sorted(output_docxs_dir.glob("*.docx"))
    if not docx_files:
        return

    # A throwaway user profile avoids LibreOffice's single-instance lock (so it
    # works even if the user already has LibreOffice open).
    profile_uri = Path(tempfile.mkdtemp(prefix="lo_levitool_")).as_uri()
    logger.debug(f"libreoffice: converting {len(docx_files)} docx files via {soffice}")
    logger.info(f"Ready to convert {len(docx_files)} docx files to pdfs (status will not be shown until finished)")
    result = subprocess.run(
        [soffice, f"-env:UserInstallation={profile_uri}", "--headless",
         "--convert-to", "pdf", "--outdir", str(output_pdfs_dir),
         *[str(f) for f in docx_files]],
        capture_output=True, text=True)
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice conversion failed:\n{result.stderr.strip()}")

    for docx_file in docx_files:
        expected_pdf = output_pdfs_dir / (docx_file.stem + ".pdf")
        if not expected_pdf.exists():
            raise RuntimeError(
                f"LibreOffice did not produce {expected_pdf.name}\n"
                f"stdout: {result.stdout.strip()}\nstderr: {result.stderr.strip()}")
    logger.info(f"libreoffice converted {len(docx_files)} docx files to PDF")


def levitool() -> None:
    """Runs the full LeviTool mail-merge pipeline: select inputs, merge data, output docx and PDF files.

    Shows a warning popup if any TEST flags are active, then calls each pipeline
    stage in sequence: directory selection, file loading, key extraction, Excel
    validation, data filtering, output directory setup, document generation, and
    PDF conversion.
    """

    _test_active = [name for name, val in [('TEST_INPUT_DIR', TEST_INPUT_DIR),
                                       ('TEST_SKIP_PROMPTS', TEST_SKIP_PROMPTS)] if val]
    if _test_active:
        pyautobek.alert("TEST FLAGS ARE SET:\n\n" + "\n".join(_test_active), "WARNING: Test Mode Active")

    path_of_input, path_root = select_input_directory(INITIAL_ATTACHMENT_DIR)
    boe_input_xlsx, docx_input_files = load_template_files(path_of_input)
    required_keys = extract_required_keys(docx_input_files)
    counties_df = load_and_validate_excel(boe_input_xlsx, required_keys)
    used_counties_df = filter_and_check_data(counties_df)
    output_docxs_path, output_pdfs_dir = setup_output_dirs(path_root)
    count_counties_processed = generate_documents(used_counties_df, docx_input_files, output_docxs_path)

    # convert created docx files to pdfs
    convert_to_pdfs(output_docxs_path, output_pdfs_dir)

    msg = f"Completed scripts and Avery sheets for {count_counties_processed} counties total."
    logger.info(msg)
    pyautobek.alert(msg, "Alert")


def main() -> None:
    """Presents a menu for selecting which LeviTool function to run.

    Handles the --update flag to reinstall the package via uvx. Otherwise displays
    a dialog for the user to choose from: LeviTool mail merge, max label length check,
    URL check, or short URL check.
    """
    import sys
    setup_loguru("INFO", "DEBUG")   # configure logging once, here (see note near imports)
    if "--update" in sys.argv:
        subprocess.run(["uvx", "--reinstall", "--from",
            "git+https://github.com/kramsman/uv_levitool.git", "levitool"])
        return

    choice = pyautobek.confirm("What do you want to do?", title='Select',
                      buttons=['LeviTool', 'Max Row Len', 'Check Urls', 'Check Short Urls',
                               'Edit Config', 'Exit'])
    choice = choice.lower()

    if choice == 'levitool':
        levitool()
    elif choice == 'max row len':
        max_label_lengths(initial_attachment_dir=INITIAL_ATTACHMENT_DIR)
    elif choice == 'check urls':
        check_boe_urls(key_containing_url=None, key_for_used=None, open_browser=True)
    elif choice == 'check short urls':
        check_short_boe_urls(key_containing_short_url=None, key_containing_url=None, key_for_used=None,
                             open_browser=None)
    elif choice == 'edit config':
        from .edit_config import edit_config
        edit_config()
    else:
        exit()


if __name__ == '__main__':

    main()
