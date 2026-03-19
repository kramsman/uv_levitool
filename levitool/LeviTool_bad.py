"""
LeviTool_bad.py creates the attachments need for postcard writing namely scripts and labels.  It does this by performing
a "mail merge" on input files, replacing fields in {} with those from an excel input file.
All the input files must be in a directory named 'Input', and must contain one xlsx with the info to be merged and
any number of docx files which will have the dat merged into them.
A pdf of each docs is created after the merge.
A separate subdirectory under a min directory named "Output" is created for each used row in the input xlsx file.
"""

# 1/31/24 Renamed LeviToolWin5.0.py to LeviTool_bad.py and started git
# done 2/1/24 strip values to be replaced.  Tab or pads on phone or county creates off center labels and could run over.
#   sb fixed by reading in county sheet w astype(str): errors if val is not a string (eg blank = nan for {url})
#   stripped vals to remove unseen chars on phone/county when reading in county sheet


# TODO currently blanks replaced with "' '" in df used_counties_df.  OK or should be blank? Or '.'?

import subprocess
# uv add uvbekutils --upgrade-package uvbekutils
subprocess.run(["uv", "add", "git+https://github.com/kramsman/uvbekutils.git"])
# subprocess.run(["uv", "add", "uvbekutils --upgrade-package uvbekutils"])
# uv add

from uvbekutils import select_file
from uvbekutils import pyautobek
# from uvbekutils import exit_yes_no
from docx import Document  # package in Conda is python-docx, not simply docx
from docx2pdf import convert
from copy import deepcopy
import os
import re
import sys
import time
import pathlib
from pathlib import Path
import glob
import inspect
import pandas as pd
# import pymsgbox
# import PySimpleGUI as sg
from loguru import logger
from collections import Counter

log_level = "DEBUG"  # used for log file; screen set to INFO. TRACE, DEBUG, INFO, WARNING, ERROR

INITIAL_ATTACHMENT_DIR = pathlib.Path("~/Dropbox/Postcard Files/Attachments/Campaigns/").expanduser()

# Fields to report max length for label checking
length_check_fields = ['CNTYFILENAME', 'CNTYTOPRINT', 'PHONE']
# field to use when checking only rows we are writing
length_check_fields_select = 'PRIORITY'


def exit_yes_no(msg, title=None, display_exiting=False):
    """ makes this choice to continue one line"""
    if not title:
        title = "Exit?"
    choice = pyautobek.confirm(msg, title, ['Yes', 'No'])
    if choice == "No":
        if display_exiting:
            pyautobek.alert("Exiting", "Alert")
        logger.error("Exiting'")
        exit()


def set_logfile_path(log_path):
    """ set log file path to location based on whether setup file is used or not.  If not, use EXE.
        using loguru could put in downloads (next line)
        but I'm locating it once we know what section
        is running, either with the setup file or the EXE if no setup, eg updating zip file.
        logfile = pathlib.Path.home() / "Downloads" / (Path(__file__).name + ".log")
    """

    logger.remove(0)
    logfile = log_path / (Path(__file__).name + ".log")
    try:
        os.remove(logfile)
    except Exception as e:
        logger.exception(e)
        pass

    logger.add(open(logfile, 'w'), level=log_level, backtrace=True, diagnose=True)
    logger.add(sys.stdout, level='INFO', backtrace=True, diagnose=True)


def calling_func(level=0):
    """ returns the various levels of calling function.  0 is current, 1 is caller of current, etc """
    try:
        func = f"'{inspect.stack()[level][3]}', line #: {inspect.stack()[level][2]}"
    except Exception as e:
        logger.exception(e)
        func = f"** error ** inspect level too deep: {str(level)} called from {inspect.stack()[level][3]}"
    return func


def exit_yes(msg: str, title: str = None, *, errmsg: str = None) -> None:
    """ exits program after giving user a popup window and raising an error. """
    msg = (msg + "\n\n\nExiting." +
           f"\n\nCalled from {calling_func(level=3)}"
           f"\nCalled from {calling_func(level=2)}"
           f"\nCalled from {calling_func(level=1)}"
           )
    if not errmsg:
        errmsg = msg.replace("\n", " ")  # dont fill the console with linefeeds
    if not title:
        title = "** Exiting Program **"
    logger.error("exiting")
    pyautobek.alert(msg, title)
    raise Exception(errmsg)


# def get_dir_name(box_title, title2, initial_dir):
#     """ show an "Open" dialog box and return the selected directory. Replaced askdirectory with PySimpleGUI
#     :param title2:
#     :type title2:
#     """
#     logger.debug('here')
#
#     layout = [
#         [sg.Text(title2, font=("Arial", 18))],
#         [
#          sg.Input(key="-IN-", expand_x=True),
#          sg.FolderBrowse(initial_folder=os.path.expanduser(initial_dir))
#          ],
#         [sg.Button("Choose")],
#     ]
#
#     # event, values = sg.Window(heading_in_box, layout, size=(600, 100)).read(close=True)
#     event, values = sg.Window(box_title, layout, titlebar_font=("Arial", 20), font=("Arial", 14),
#                               size=(1000, 150), use_custom_titlebar=True).read(close=True)
#
#     dir_name = values['-IN-']
#     if dir_name == "":
#         exit_yes("No directory name chosen")
#
#     return dir_name


# def format_ok(s):
#
#     # more pythonic way
#     # FIXME shouldn't all non 0 be false?  why the separate check returning False?
#     # bracket_count = s.count('{') - s.count('}')
#     # if bracket_count < 0 or bracket_count > 1:
#     #     bracket_chk = False
#     # elif bracket_count == 0:
#     #     bracket_chk = True
#     # else:
#     #     bracket_chk = True
#
#     bracket_count = 0
#     for char in s:
#         if char == '{':
#             bracket_count += 1
#         elif char == '}':
#             bracket_count -= 1
#
#         if bracket_count < 0 or bracket_count > 1:
#             return False
#
#     return bracket_count == 0


def count_brackets(s):
    """ counts the difference in the number of left and right brackets.
    returns -10000 if difference is more than 1 """

    # more pythonic way
    bracket_count = s.count('{') - s.count('}')
    if bracket_count < -1 or bracket_count > 1:
        bracket_count = -10000

    # bracket_count = 0
    # for char in s:
    #     if char == '{':
    #         bracket_count += 1
    #     elif char == '}':
    #         bracket_count -= 1
    #
    #     if bracket_count < -1 or bracket_count > 1:
    #         return -10000

    return bracket_count


def find_keys(document):
    # Keys are fields in document to be replaced
    keys = set()

    d = deepcopy(document)

    for section in d.sections:
        for paragraph in section.header.paragraphs:
            new_keysx = re.findall(r'{[a-zA-Z0-9]+}', paragraph.text)
            new_keys = [x.upper() for x in new_keysx]
            keys = keys.union({k.upper() for k in new_keys})

    for paragraph in d.paragraphs:
        new_keysx = re.findall(r'{[a-zA-Z0-9]+}', paragraph.text)
        new_keys = [x.upper() for x in new_keysx]
        keys = keys.union({k.upper() for k in new_keys})

    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    new_keysx = re.findall(r'{[a-zA-Z0-9]+}', paragraph.text)
                    new_keys = [x.upper() for x in new_keysx]
                    keys = keys.union({k.upper() for k in new_keys})

    return keys


def replace_tags(paragraph, county_row):
    brackets = 0
    partial_runs = []
    for run in paragraph.runs:
        partial_runs.append(run)
        brackets += count_brackets(run.text)
        if brackets == 0:
            text = ''.join([r.text for r in partial_runs])
            for key, val in county_row.items():
                # print(f'{key=}, {val=}')
                text = re.sub(key, val, text, flags=re.IGNORECASE)  # use re to replace and ignore case
            partial_runs[0].text = text
            for r in partial_runs[1:]:
                r.text = ''
            
            brackets = 0
            partial_runs = []
        elif brackets != 1:
            raise ValueError(f'Template error near "{run.text}" in script template. Make sure sequence to be replaced has consistent formatting.')


def make_replacements(document, county_row):

    for paragraph in document.paragraphs:
        replace_tags(paragraph, county_row)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_tags(paragraph, county_row)

    return


def main():

    def check_header_count(hdr_str, header_count, num_allowed):
        """ make sure headers appear in file only given number of times and error if not"""

        if header_count != num_allowed:
            msg = f"'{hdr_str}' does not appear {num_allowed} time(s), instead {header_count}\n\nEXITING."
            # logger.error(msg)
            # messagebox.showerror("ERROR", msg)
            # exit()
            exit_yes(msg, "ERROR")

    def prompt_for_df_values(df, msg, box_title):
        """ display values of a dataframe after converting to string and prompt to continue"""

        df_string = df.to_string(index=False)

        choice = pyautobek.confirm(f"{msg}\n\n"
                                  f"{df_string}\n\n",
                                  box_title,
                                  ['Continue', 'Exit'])
        if choice == "Exit":
            logger.info("Exiting")
            exit()

    # Select input dir: True via dialog; False hardcoded
    if False:

        # path_to_add = '/Users/Denise/Library/CloudStorage/Dropbox/non ROV python progs/uvcleaver'
        # sys.path.append(str(Path(path_to_add).resolve()))
        # from select_file import select_file  # type: ignore  # tells the linter to not flag missing library that
        # will be added during
        # run

        # title: str, start_dir: str, files_like: str, choices: list[str], mode: str = "file"

        # path_input_str = select_file("Pick 'Input' Directory",
        #                              INITIAL_ATTACHMENT_DIR,
        #                               # "Select the directory containing the input files for LeviTool (eg .../GA files/Input')",
        #                              '',
        #                               mode = 'dir',
        #                              )
        path_input_str = select_file(
            title="Pick 'Input' Directory",
            start_dir=INITIAL_ATTACHMENT_DIR,
            files_like="input",
            choices=["Select", "Cancel"],
            mode="dir"  # file, dir or both
        )
        # path_input_str = get_dir_name("Pick 'Input' Directory",
        #                               "Select the directory containing the input files for LeviTool (eg .../GA "
        #                               "files/Input')",
        #                               INITIAL_ATTACHMENT_DIR)

        path_input = pathlib.Path(path_input_str)
    else:
        path_input = pathlib.Path("~/Dropbox/Postcard Files/Attachments/Campaigns/Test box and image/Input").expanduser()
    path_root = path_input.parents[0]

    if path_input.parts[-1].upper() != "INPUT":
        msg = f"Last part of chosen directory:\n\n {path_input}\n\nDoes not equal 'INPUT'.\n\nEXITING."
        exit_yes(msg, "ERROR")

    # list of all xlsx files not starting with ~ (temporary files)
    xlsx_files = glob.glob(str(path_input / '[!~]*.xlsx'))

    if len(xlsx_files) != 1:
        msg = f"Must have only 1 xlsx file in the Input directory; there are {len(xlsx_files)}.  EXITING."
        exit_yes(msg, "ERROR")

    script_file_name = xlsx_files[0]

    # Get list of all template files (docx) that we will update not starting with ~ (temporary files).
    docx_files = glob.glob(str(path_input / '[!~]*.docx'))

    if len(docx_files) < 1:
        # print('ERROR: No docx template found in directory: ' + path_input + 'EXITING.')
        msg = f"No docx input file found in directory: '{path_input}' EXITING."
        exit_yes(msg, "ERROR")

    logger.info(f"\nWill replace tags in {len(docx_files)} documents per county.\n")

    # loop through each docx template file and put keys in required_keys set
    required_keys = set()
    for docx_file in docx_files:
        document_of_docx_file = Document(docx_file)
        keys_in_doc = find_keys(document_of_docx_file)

        required_keys = required_keys.union(keys_in_doc)

    # Add the keys required by the program but not necessarily used for replacement in docxs.
    required_keys.add('{STATE}')
    required_keys.add('{CNTYFILENAME}')  # added BEK 4/23/22
    required_keys.add('{USE}')
    # add keys for those we are checking max length
    for field in length_check_fields:
        required_keys.add(f"{{{field}}}")
    # add field used in filtering rows for max length check
    if length_check_fields_select:
        required_keys.add(f"{{{length_check_fields_select}}}")

    # Note .astype(str) below which reads all data as string and converts missing to 'nan'
    county_sheet_df = pd.read_excel(script_file_name, sheet_name='Counties', header=None, skiprows=1,
                                    nrows=1).astype(str)
    # read only row 2 (1 is skipped above) to get keys
    county_sheet_keys = county_sheet_df.loc[0, :].values.tolist()
    county_sheet_keys = [x.upper() for x in county_sheet_keys if x != "nan"]

    # check that three key tokens only appear once (needed and may not be on file)
    check_header_count('{USE}', county_sheet_keys.count('{USE}'), 1)
    check_header_count('{STATE}', county_sheet_keys.count('{STATE}'), 1)
    check_header_count('{CNTYFILENAME}', county_sheet_keys.count('{CNTYFILENAME}'), 1)

    # check for one occurrence for keys for those we are checking max length
    for field in length_check_fields:
        check_header_count(f"{{{field}}}", county_sheet_keys.count(f"{{{field}}}"), 1)

    # check for one occurrence of field if we are checking length for a subset we are writing
    if length_check_fields_select:
        check_header_count(f"{{{length_check_fields_select}}}",
                           county_sheet_keys.count(f"{{{length_check_fields_select}}}"), 1)

    # keys in file heading also being replaced in docxs
    county_sheet_required_keys = [c.upper() for c in county_sheet_keys if c.upper() in required_keys]

    # get key values and number of occurrences
    key_counts = Counter(county_sheet_required_keys)

    bad_key_counts = [(k, key_counts[k]) for k in key_counts if key_counts[k] > 1]
    if bad_key_counts:
        fld_string = '\n'.join([f"{tup[0]}: {str(tup[1])}" for tup in bad_key_counts])
        msg = f"The following column keys are in the sheet more than once:\n\n{fld_string}"
        exit_yes(msg, "Duplicate Key Columns in Sheet")

    county_sheet = pd.read_excel(script_file_name, sheet_name='Counties', header=0, skiprows=1).astype(str)
    county_sheet.columns = [c.upper() for c in county_sheet.columns]  # convert all column names to uppercase
    # get list of required columns
    required_county_cols = [c.upper() for c in county_sheet.columns if c.upper() in required_keys]

    # only keep needed columns in df
    county_sheet = county_sheet[required_county_cols]
    # replace fields with stripped values then "' '" to account for and see multi spaces
    county_sheet = county_sheet.apply(lambda x: x.str.strip())

    # get max lengths of county to check label
    for field in length_check_fields:
        max_field = max(county_sheet[f"{{{field}}}"], key=len)
        logger.info(f"'{max_field}': Max of '{{{field}}}' column")

    # get max lengths of county to check label for rows selected
    if length_check_fields_select:
        for field in length_check_fields:
            # max_field = max(county_sheet.loc[county_sheet[f"{{{length_check_fields_select}}}"] != 'nan']
            max_field = max(county_sheet.loc[county_sheet[f"{{{length_check_fields_select}}}"].notnull()]
                                             [f"{{{field}}}"], key=len)
            logger.info(f"'{max_field}': Max of '{{{field}}}' where '{{{length_check_fields_select}}}' is not blank")

    import pyautogui

    # Creating a choice box with custom options
    choice = pyautogui.confirm("Do you want to continue running or set {USE} based on the longest substitutions? Check the log for values.\n\nContinue?",
                               title="Check the log for values.\n\nContinue?", buttons=['Yes', 'No'])

    if choice == "No":
        # if display_exiting:
        #     pyautobek.alert("Exiting", "Alert")
        logger.debug("here")
        exit()
    else:
       print("Continuing")

    # exit_yes_no("Do you want to continue running or set {USE} based on the longest substitutions?  "
    #             "Check the log for values.\n\nContinue?",
    #             "Continue or Set {USE} by Field Max Lengths?")

    # Create df of ony those counties being processed.  Used below.  Fields are cleaned up here.
    # used_counties_df = county_sheet.loc[county_sheet['{USE}'] != 'nan']
    used_counties_df = county_sheet.loc[county_sheet['{USE}'].notnull()]
    # xxx = county_sheet.loc[county_sheet['{USE}'] == 'nan']
    xxx = county_sheet.loc[county_sheet['{USE}'].isnull()]

    # replace fields with string showing quotes ("' '") to make easier to see.
    # May not want to replace in file used for replacements so we could replace keys with spaces
    # in docxs
    used_counties_df.replace([""], ["' '"], inplace=True)

    # check if any fields to be substituted have nan values
    # nan_in_df = used_counties_df[used_counties_df.eq('nan').any(axis=1)]
    nan_in_df = used_counties_df[used_counties_df.isnull().any(axis=1)]
    if not nan_in_df.empty:
        prompt_for_df_values(nan_in_df,
                             "Some records contain missing ('nan') values that will be replaced in docxs:",
                             "OK to Substitute Missing Values ('nan's)?")

    # check if any fields to be substituted have blank values
    blank_in_df = used_counties_df[used_counties_df.eq("' '").any(axis=1)]  # all fields replaced with stripped val then ' '  above to see
    if not blank_in_df.empty:
        prompt_for_df_values(blank_in_df,
                             "Some records contain missing blank values that will be replaced in docxs:",
                             "OK to Substitute Blank Values?")

    # Get number of counties in sheet with non-blank {USE}
    # num_counties_to_use = county_sheet.loc[county_sheet['{USE}'] != 'nan', '{USE}'].count()
    num_counties_to_use = len(used_counties_df)  # number of records

    list_of_used_county_names = used_counties_df['{CNTYFILENAME}'].to_list()
    string_of_used_county_names = ",".join(list_of_used_county_names)

    # Get total number of counties in sheet with non-blank {USE}
    total_counties = county_sheet['{CNTYFILENAME}'].count()
    logger.info(f"Running on counties: {string_of_used_county_names}")
    choice = pyautobek.confirm(f"Will process {num_counties_to_use} of {total_counties} total counties:\n\n{string_of_used_county_names}",
                              'Continue or Exit?',
                              ['Continue', 'Exit'])
    if choice == "Exit":
        logger.info("Exiting")
        exit()

    choice = pyautobek.confirm("Close MS Word now if open\n\n",
                              'Continue or Exit?',
                              ['Continue', 'Exit'])
    if choice == "Exit":
        exit()

    count = 0
    for _, county_row in used_counties_df.iterrows():  # uses df of only rows with used != nan from above
        state_fn = county_row['{STATE}'] + "-" + county_row['{CNTYFILENAME}']

        for docx_file in docx_files:
            document_of_docx_file = Document(docx_file)
            script = deepcopy(document_of_docx_file)
            make_replacements(script, county_row)

            output_dir = path_root / 'Output'
            if not output_dir.exists():
                output_dir.mkdir()

            state_county_dir = output_dir / state_fn  # added BEK 4/23/22 to use filename var
            if not state_county_dir.exists():
                state_county_dir.mkdir()

            file_with_state = state_fn + ' ' + pathlib.Path(docx_file).name  # use filename var to match VL
            script_filepath = state_county_dir / file_with_state
            script.save(script_filepath)

        time.sleep(2.5)  # Dropbox needs time to update or Word bombs
        convert(state_county_dir)  # this converts all docx files in the folder to pdf using docs2pdf function

        count += 1
        msg = f"Finished filling Word templates for {state_fn}. Completed {count} of {num_counties_to_use} chosen counties."
        logger.info(msg)

    msg = f"Completed scripts and Avery sheets for {count} counties total."
    logger.info(msg)
    # messagebox was throwing an unsolvable SIGSEGV error (ie in the C code) so use msgbox
    # messagebox.showinfo("Finished", msg)
    pyautobek.alert(msg, "Alert")


if __name__ == '__main__':
    # import subprocess
    # # uv add uvbekutils --upgrade-package uvbekutils
    # # subprocess.run(["uv", "add", "git+https://github.com/kramsman/uvbekutils.git"])
    # subprocess.run(["uv", "add", "uvbekutils --upgrade-package uvbekutils"])
    # # uv add

    main()
