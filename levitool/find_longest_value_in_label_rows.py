""" Check the lengths of row text in a label after Levitool substitutions to make sure it will fit.
Should be max of 34.
"""

# FIXME check if {priority} changes the number of rows used in count

import re
from copy import deepcopy
from pathlib import Path

from docx import Document  # package in Conda is python-docx, not simply docx
from uvbekutils import pyautobek
from uvbekutils import safe_str
from uvbekutils import scroll_box
from uvbekutils import list_pick
from uvbekutils import standardize_columns

from .read_boe_xls import read_boe_xls
from .constants import DEFAULT_FONT_SIZE_PT
from .label_width import (resolve_font, max_fitting_size_pt, segments_width_pt,
                          DEFAULT_FAMILY, FIT_TOLERANCE, SUGGEST_TOLERANCE)

# {KEY} replacement tokens, e.g. '{county}'.
_KEY_RE = re.compile(r'\{[a-zA-Z0-9]+\}')

# Unicode Private Use Area. Symbol fonts (the CFCG logo among them) map their glyphs here,
# so a character in this range is a picture, not text that should be resized.
_PUA_START, _PUA_END = 0xE000, 0xF8FF


def get_label_text(label_docx) -> str:
    """Reads text from the upper-left cell of the first table in a label docx file.

    Args:
        label_docx: Path or string path to the label docx template file.

    Returns:
        str: Text from the upper-left label cell with lines joined by newline characters.
    """

    # string or path?
    document_of_docx_file = Document(label_docx)
    d = deepcopy(document_of_docx_file)

    # upper left label
    lines = [line.text for line in d.tables[0].rows[0].cells[0].paragraphs]
    label_text = '\n'.join(lines)
    print(f"{label_text=}")
    return label_text


def _theme_families(document) -> dict:
    """Returns the theme's major/minor Latin font names, e.g. {'minor': 'Aptos'}.

    A run can name its font only by theme reference (w:rFonts/@w:asciiTheme='minorHAnsi'),
    in which case the real family lives in theme1.xml. Returns an empty dict when the
    theme part is missing or unreadable.
    """
    from lxml import etree

    a = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
    families = {}
    try:
        for part in document.part.package.iter_parts():
            if not str(part.partname).endswith('theme1.xml'):
                continue
            root = etree.fromstring(part.blob)
            for key, scheme in (('major', 'majorFont'), ('minor', 'minorFont')):
                latin = root.find(f'.//{a}fontScheme/{a}{scheme}/{a}latin')
                if latin is not None and latin.get('typeface'):
                    families[key] = latin.get('typeface')
            break
    except Exception:  # a malformed or absent theme must not break the fit check
        return {}
    return families


def _rfonts_family(rfonts, themes: dict):
    """Returns the Latin family named by a w:rFonts element, or None.

    Prefers the explicit w:ascii name and falls back to resolving a w:asciiTheme
    reference through the document theme.
    """
    from docx.oxml.ns import qn

    if rfonts is None:
        return None
    name = rfonts.get(qn('w:ascii'))
    if name:
        return name
    theme_ref = rfonts.get(qn('w:asciiTheme'))
    if theme_ref:
        return themes.get('major' if 'major' in theme_ref.lower() else 'minor')
    return None


def _doc_defaults(document) -> dict:
    """Returns the run formatting Word applies when a run specifies none of its own.

    Reads the w:docDefaults block in styles.xml, then lets the Normal style override it.
    Only when the file says nothing at all do we fall back to DEFAULT_FONT_SIZE_PT and
    DEFAULT_FAMILY. Reading the real defaults matters: a template whose default is Word's
    modern 11 pt would otherwise have every unsized run (the CFCG logo run among them)
    measured at 10 pt, under-reporting the line and calling it a fit when it wraps.

    Args:
        document: An open docx Document.

    Returns:
        dict: Keys 'size' (points), 'family', 'size_from_file' and 'family_from_file'
            (False when the value is our fallback rather than the document's), and
            'themes' (from _theme_families(), for resolving run-level theme references).
    """
    from docx.oxml.ns import qn

    themes = _theme_families(document)
    size_pt = None
    family = None

    rpr = document.styles.element.find(
        qn('w:docDefaults') + '/' + qn('w:rPrDefault') + '/' + qn('w:rPr'))
    if rpr is not None:
        sz = rpr.find(qn('w:sz'))
        if sz is not None and sz.get(qn('w:val')) is not None:
            size_pt = int(sz.get(qn('w:val'))) / 2  # half-points -> points
        family = _rfonts_family(rpr.find(qn('w:rFonts')), themes)

    try:
        normal = document.styles['Normal']
    except KeyError:
        normal = None
    if normal is not None:
        if normal.font.size is not None:
            size_pt = normal.font.size.pt
        normal_family = _rfonts_family(
            normal.element.find(qn('w:rPr') + '/' + qn('w:rFonts')), themes)
        if normal_family:
            family = normal_family

    return {
        'size': size_pt if size_pt else DEFAULT_FONT_SIZE_PT,
        'family': family or DEFAULT_FAMILY,
        'size_from_file': size_pt is not None,
        'family_from_file': family is not None,
        'themes': themes,
    }


def _run_format(run, para, defaults: dict) -> dict:
    """Returns the family, size, and weight actually in force for one run.

    Each attribute is resolved through Word's inheritance chain - the run itself, then its
    character style, then the paragraph style, then the document defaults. Formatting is
    read per run rather than per line because a single label line mixes them: the CFCG
    logo run is a different family and size from the sentence beside it.

    Args:
        run: A docx Run.
        para: The paragraph containing it (for the paragraph style).
        defaults (dict): From _doc_defaults().

    Returns:
        dict: Keys 'family', 'size', 'bold', 'estimated' (size fell through to our
            constant) and 'family_assumed' (family fell through to our constant).
    """
    from docx.oxml.ns import qn

    rpr = run._r.find(qn('w:rPr'))
    char_style = run.style          # the run's character style, or the default one
    para_style = para.style

    size_pt = run.font.size.pt if run.font.size is not None else None
    if size_pt is None and char_style is not None and char_style.font.size is not None:
        size_pt = char_style.font.size.pt
    if size_pt is None and para_style is not None and para_style.font.size is not None:
        size_pt = para_style.font.size.pt
    estimated = False
    if size_pt is None:
        size_pt = defaults['size']
        estimated = not defaults['size_from_file']

    # Read w:rFonts directly rather than run.font.name so a theme reference resolves too.
    family = _rfonts_family(rpr.find(qn('w:rFonts')) if rpr is not None else None,
                            defaults['themes'])
    if family is None and char_style is not None:
        family = char_style.font.name
    if family is None and para_style is not None:
        family = para_style.font.name
    family_assumed = False
    if family is None:
        family = defaults['family']
        family_assumed = not defaults['family_from_file']

    bold = run.font.bold
    if bold is None and char_style is not None:
        bold = char_style.font.bold
    if bold is None and para_style is not None:
        bold = para_style.font.bold

    return {'family': family, 'size': size_pt, 'bold': bool(bold),
            'estimated': estimated, 'family_assumed': family_assumed}


def _has_pua(text: str) -> bool:
    """True when text contains a Private Use Area character (a symbol-font glyph)."""
    return any(_PUA_START <= ord(c) <= _PUA_END for c in text)


def _mark_symbol_segments(segments) -> None:
    """Flags glyph segments as symbols, in place.

    Segments already flagged - from a <w:sym> element or containing a Private Use Area
    character - are left alone. This catches the remaining case: a glyph typed as an
    ordinary character in a symbol font. The test is deliberately narrow (at most two
    visible characters, in a family other than the line's main one) so a legitimately
    differently-fonted *word* is still treated as text that can be resized.
    """
    counts = {}
    for seg in segments:
        if seg['is_symbol']:
            continue
        counts[seg['family']] = counts.get(seg['family'], 0) + len(seg['text'])
    if not counts:
        return
    dominant = max(counts, key=counts.get)
    for seg in segments:
        if not seg['is_symbol'] and seg['family'] != dominant and len(seg['text'].strip()) <= 2:
            seg['is_symbol'] = True


def _para_segments(para, defaults: dict) -> list:
    """Splits a paragraph into formatting segments in document order.

    Walks the paragraph's inner content so runs nested in a <w:hyperlink> are included -
    python-docx omits those from paragraph.runs while still reporting their text, which
    left a hyperlinked URL measured with formatting invented from its neighbours. Within a
    run, <w:sym> elements are turned into their character: Word's Insert > Symbol writes no
    w:t at all, so those glyphs previously contributed zero width.

    Args:
        para: A docx Paragraph.
        defaults (dict): From _doc_defaults().

    Returns:
        list: Segment dicts with 'text', 'family', 'size', 'bold', 'is_symbol', and the
            'estimated'/'family_assumed' flags from _run_format().
    """
    from docx.oxml.ns import qn
    from docx.text.hyperlink import Hyperlink

    segments = []
    for item in para.iter_inner_content():
        runs = item.runs if isinstance(item, Hyperlink) else [item]
        for run in runs:
            fmt = _run_format(run, para, defaults)
            for child in run._r:
                sym_family = None
                if child.tag == qn('w:t'):
                    text = child.text or ''
                elif child.tag == qn('w:tab'):
                    text = '\t'
                elif child.tag == qn('w:sym'):
                    char = child.get(qn('w:char'))
                    if not char:
                        continue
                    text = chr(int(char, 16))
                    sym_family = child.get(qn('w:font'))
                else:
                    continue  # w:br, w:drawing, proofing marks: no text advance to add
                if not text:
                    continue
                seg = dict(fmt)
                seg['text'] = text
                seg['is_symbol'] = sym_family is not None or _has_pua(text)
                if sym_family:
                    seg['family'] = sym_family
                segments.append(seg)

    _mark_symbol_segments(segments)
    return segments


def segments_text(segments) -> str:
    """Returns the plain text of a line's segments."""
    return ''.join(seg['text'] for seg in segments)


def _strip_trailing(segments) -> list:
    """Copy of segments with trailing whitespace dropped.

    Word ignores trailing spaces when deciding where to wrap, so they must not count
    toward the measured width. Leading spaces are kept - those Word does render.
    """
    out = [dict(seg) for seg in segments]
    while out:
        stripped = out[-1]['text'].rstrip()
        if stripped:
            out[-1]['text'] = stripped
            break
        out.pop()
    return out


def substitute_segments(segments, values: dict) -> list:
    """Replaces {KEY} tokens in a line's segments, keeping each piece's formatting.

    Word splits a single token across runs - one label line is stored as 'For {' /
    'cntytoprint' / '} info:' - so substitution cannot be done run by run. The segments are
    joined into one string alongside a record of which segment each character came from;
    tokens are matched against that joined text and each replacement value inherits the
    formatting of the token's first character.

    Args:
        segments (list): Template segments for one label line, from _para_segments().
        values (dict): Lowercased '{key}' -> replacement string.

    Returns:
        list: New segment dicts for the substituted line.
    """
    text = segments_text(segments)
    owner = []  # segment index of each character in text
    for i, seg in enumerate(segments):
        owner.extend([i] * len(seg['text']))

    pieces = []  # (segment index, text)

    def add_span(start, end):
        """Copies text[start:end] through, split at each segment boundary it crosses."""
        i = start
        while i < end:
            j = i
            while j < end and owner[j] == owner[i]:
                j += 1
            pieces.append((owner[i], text[i:j]))
            i = j

    last = 0
    for match in _KEY_RE.finditer(text):
        start, end = match.span()
        add_span(last, start)
        key = match.group(0).lower()
        if key in values:
            pieces.append((owner[start], safe_str(values[key])))
        else:
            add_span(start, end)  # no column for this token; leave it visible
        last = end
    add_span(last, len(text))

    out = []
    for seg_index, piece_text in pieces:
        if not piece_text:
            continue
        if out and out[-1]['_from'] == seg_index:
            out[-1]['text'] += piece_text
            continue
        seg = dict(segments[seg_index])
        seg['text'] = piece_text
        seg['_from'] = seg_index
        out.append(seg)
    for seg in out:
        del seg['_from']
    return out


def _para_indent_pt(para) -> float:
    """Returns the total horizontal indent (points) that narrows a paragraph's text.

    Word reduces a paragraph's usable text width by its left and right indents, and the
    first line additionally by a positive first-line indent. A label line is a single
    visual line, so the first-line indent (when positive) applies to it. These indents
    are on top of the cell margins and must be subtracted from the cell width, or the
    fit check overestimates how much text fits (which let a wrapping line slip through).
    """
    pf = para.paragraph_format
    left = pf.left_indent.pt if pf.left_indent is not None else 0.0
    right = pf.right_indent.pt if pf.right_indent is not None else 0.0
    first = pf.first_line_indent.pt if pf.first_line_indent is not None else 0.0
    return left + right + max(first, 0.0)


def get_label_line_segments(label_docx) -> list:
    """Returns the label cell's lines, each split into formatting segments.

    Paragraphs that are blank in the template are dropped here, and each line's formatting
    travels with its text from this point on. Previously the font info was indexed by
    unfiltered paragraph number while blank lines were dropped from the text list, so a
    single blank paragraph in the label cell silently shifted every later line's font and
    size onto the wrong line. A line that empties only after substitution is kept - it
    simply measures as nothing - which also keeps every county row the same length.

    Args:
        label_docx: Path or string path to the label docx template file.

    Returns:
        list: One dict per non-blank line, in order, with keys 'line' (1-based number),
            'indent_pt' (horizontal paragraph indent in points, which narrows this line's
            usable width), and 'segments' (see _para_segments()).
    """
    d = deepcopy(Document(label_docx))
    defaults = _doc_defaults(d)
    lines = []
    for para in d.tables[0].rows[0].cells[0].paragraphs:
        segments = _para_segments(para, defaults)
        if not segments_text(segments).strip():
            continue
        lines.append({
            'line': len(lines) + 1,
            'indent_pt': _para_indent_pt(para),
            'segments': segments,
        })
    return lines


def _format_description(segments) -> str:
    """Human-readable summary of the sizes, weights, and families used on a line.

    Returns something like '9.5 pt bold Arial + 10 pt CFCG Symbol' so a mixed-format line
    is visible in the report rather than hidden behind one made-up size.
    """
    seen = []
    for seg in segments:
        if not seg['text'].strip() or not seg['size']:
            continue
        family_used = resolve_font(seg['family'], seg['bold'])[1]
        desc = f"{seg['size']:g} pt{' bold' if seg['bold'] else ''} {family_used}"
        if desc not in seen:
            seen.append(desc)
    return ' + '.join(seen)


def _format_notes(segments) -> str:
    """Parenthesised note listing anything on the line that had to be assumed."""
    notes = []
    if any(seg['estimated'] for seg in segments):
        notes.append("size est.")
    assumed = []
    for seg in segments:
        _, family_used, fallback = resolve_font(seg['family'], seg['bold'])
        if (fallback or seg['family_assumed']) and family_used not in assumed:
            assumed.append(family_used)
    if assumed:
        notes.append(f"font assumed {', '.join(assumed)}")
    return f" ({', '.join(notes)})" if notes else ""


def get_label_cell_usable_width_pt(label_docx) -> float:
    """Returns the usable text width (in points) of the label's upper-left cell.

    Usable width is the cell width minus its left and right margins. The cell width is
    read from w:tcW (falling back to the table grid column), and margins from the cell's
    own w:tcMar, then the table default w:tblCellMar; when no margin is specified,
    Word's default of 108 twips (0.075 in) per side is assumed. Returns None if the cell
    width cannot be determined. Twips are converted at 20 twips per point.

    Args:
        label_docx: Path or string path to the label docx template file.

    Returns:
        float: Usable cell width in points, or None if it cannot be determined.
    """
    from docx.oxml.ns import qn

    # Word's default table cell side margin when none is specified.
    DEFAULT_SIDE_MARGIN_TWIPS = 108

    d = deepcopy(Document(label_docx))
    tbl = d.tables[0]
    cell = tbl.rows[0].cells[0]
    tcPr = cell._tc.find(qn('w:tcPr'))

    def _w(element):
        """Reads a w:w attribute (twips) off an element, or None."""
        if element is None:
            return None
        val = element.get(qn('w:w'))
        return int(val) if val is not None else None

    # Cell width: prefer the cell's own w:tcW, else the first grid column.
    cell_w = _w(tcPr.find(qn('w:tcW'))) if tcPr is not None else None
    if cell_w is None:
        grid_col = tbl._tbl.find(qn('w:tblGrid') + '/' + qn('w:gridCol'))
        cell_w = _w(grid_col)
    if cell_w is None:
        return None

    def _side_margin(container, mar_tag, side):
        """Reads a left/right margin (twips) from a *Mar element, or None."""
        if container is None:
            return None
        mar = container.find(qn(mar_tag))
        if mar is None:
            return None
        return _w(mar.find(qn('w:' + side)))

    tblPr = tbl._tbl.tblPr
    left = _side_margin(tcPr, 'w:tcMar', 'left')
    if left is None:
        left = _side_margin(tblPr, 'w:tblCellMar', 'left')
    if left is None:
        left = DEFAULT_SIDE_MARGIN_TWIPS
    right = _side_margin(tcPr, 'w:tcMar', 'right')
    if right is None:
        right = _side_margin(tblPr, 'w:tblCellMar', 'right')
    if right is None:
        right = DEFAULT_SIDE_MARGIN_TWIPS

    usable_twips = cell_w - left - right
    return usable_twips / 20.0  # 20 twips per point


def max_label_lengths(*, used_field: str = None, label_docx=None, initial_attachment_dir=None) -> None:
    """Checks the maximum line length in a label template after mail-merge substitutions.

    Prompts for the label docx and auto-detects the BOE xlsx from the same directory.
    Substitutes all {KEY} tokens for every county row, then reports the widest line
    per label line number, measuring each line's real rendered width (actual font,
    size, and weight) against the label cell's usable width. Displays results and any
    overflow warnings in a scroll box.

    Args:
        used_field (str): Column key identifying active county rows (e.g. '{priority}').
            Prompts if None.
        label_docx: Path to the label docx template. Prompts via file picker if None.
        initial_attachment_dir: Starting directory for the label docx file picker.
            Used only when label_docx is None.
    """

    import pandas as pd
    from pathlib import Path
    from loguru import logger
    import glob
    from uvbekutils import exit_yes, select_file

    def print_max_line_info(df: pd.DataFrame, line_list_field: str, line_meta: list = None,
                            usable_width_pt: float = None):
        """Finds and prints the widest substituted line for each line position in the label.

        Measures each substituted line segment by segment - in each segment's own font
        family, point size, and weight - then compares the total against the label cell's
        usable width and appends a warning (with a suggested smaller size) if the line will
        not fit. The widest line is chosen by measured width, not character count, because
        fonts are proportional.

        Args:
            df (pd.DataFrame): DataFrame where each row contains a list of substituted
                label lines (each itself a list of segments) in line_list_field.
            line_list_field (str): Name of the column containing per-row lists of label lines.
            line_meta (list): Optional per-line dicts from get_label_line_segments(),
                supplying each line's paragraph indent.
            usable_width_pt (float): Usable cell width in points from
                get_label_cell_usable_width_pt(); the fit limit for every line.

        Returns:
            tuple: (summary, has_overflow). summary is a newline-joined report of the
                widest line, its sizes/fonts/width, and any overflow warning for each label
                line position, or 'None' if empty. has_overflow is True when at least one
                line will not fit.
        """
        has_overflow = False
        if df.empty:
            pyautobek.alert(f"No counties are being selected based on the field '{used_field}'.\n\n"
                            f"Checking for these copunties will be skipped.\n"
                            f"If desired, rerun choosing another field (like {{priority}}'" ,
                        "No Counties Being Selected")
            result = 'None'
        else:
            # get number of list elements (# of lines) in the label, so we know how many need to be checked
            number_of_lines = len(df[line_list_field].iloc[0])
            result_lines = []
            for line_number in range(number_of_lines):
                # Trailing spaces do not count toward the width - Word ignores them when
                # deciding where to wrap.
                line_data = [_strip_trailing(line[line_number]) for line in df[line_list_field]]

                info = line_meta[line_number] if line_meta and line_number < len(line_meta) else {}

                # This line's usable width = the cell's usable width minus this
                # paragraph's own indent (indents narrow the text and vary per line).
                line_usable = None
                if usable_width_pt is not None:
                    line_usable = usable_width_pt - info.get('indent_pt', 0.0)

                # Measure every candidate line and keep the widest by RENDERED width (not
                # char count, since fonts are proportional). A width is None only when a
                # segment has no size or no usable font file on this machine.
                widths = [(segs, segments_width_pt(segs)) for segs in line_data]
                measurable = [(segs, w) for segs, w in widths if w is not None]
                if measurable:
                    max_segments, measured = max(measurable, key=lambda t: t[1])
                else:
                    max_segments = max(line_data, key=lambda segs: len(segments_text(segs)))
                    measured = None
                max_line = segments_text(max_segments)
                note_str = _format_notes(max_segments)

                if measured is not None and line_usable is not None and line_usable > 0:
                    # A guard band, because Pillow and Word do not agree to the last
                    # fraction of a point and a real overflow can be under half a point.
                    over = measured > line_usable * FIT_TOLERANCE
                    meta_str = (f"  - {_format_description(max_segments)}, "
                                f"{measured:.0f} of {line_usable:.0f} pt wide{note_str}")
                    if over:
                        has_overflow = True
                        suggested = max_fitting_size_pt(max_segments, line_usable)
                        if suggested and suggested >= 5:
                            fit_str = f"try {suggested:g} pt"
                        else:
                            fit_str = "shorten text - won't fit"
                        warning = f"   - **** TOO WIDE - {fit_str} ****"
                    else:
                        warning = "   - fits"
                else:
                    # No font metrics available (font size unknown or no font found):
                    # report the longest line by character count without a verdict.
                    meta_str = f"  - {len(max_line)} chars (width not measured{note_str})"
                    warning = ""

                line_info = f"Line {line_number + 1},  '{max_line}'{meta_str}{warning}"
                print(line_info)
                result_lines.append(line_info)
            result = '\n'.join(result_lines)

        return result, has_overflow

    if label_docx is None:
        label_docx = select_file("PICK LABEL DOCX",
                                 start_dir=initial_attachment_dir,
                                 files_like='*.docx',
                                 mode='file',
                                 title2="Pick label docx template that will have BOE info merged into it "
                                        "('LABEL 30 per page...')",
                                 )
    logger.info(f"Label docx being used for input: '{label_docx}")
    label_docx = Path(label_docx).expanduser()

    label_text = get_label_text(label_docx)
    label_lines = get_label_line_segments(label_docx)
    usable_width_pt = get_label_cell_usable_width_pt(label_docx)
    logger.info(f"Label cell usable width: {usable_width_pt} pt")
    keys = find_keys_in_text(label_text)
    keys = [key.lower() for key in keys]

    # list of all xlsx files not starting with ~ (temporary files)
    xlsx_files = glob.glob(str(label_docx.parent / '[!~]*.xlsx'))

    if len(xlsx_files) != 1:  # must have 1, not more
        msg = f"Can not select BOE xlsx file from label directory because there are more than one;" \
              f"there are {len(xlsx_files)}.  EXITING."
        exit_yes(msg, "MORE THAN ONE XLSX")
    # xlsx_file_name = xlsx_files[0]  # filename of the first (and only) xlsx
    boe_xls = Path(xlsx_files[0]).expanduser()

    county_sheet_keys, df = read_boe_xls(boe_xls)

    # Lowercase all column names to match lowercase keys
    df.columns = [col.lower() for col in df.columns]

    if used_field is None:
        used_field = list_pick(lst=df.columns,
                               title="PICK KEY FIELD TO ID 'USED' COUNTIES",
                               msg="All counties with this field not empty will be checked as a group ("
                                   "'{priority}' will be used if none chosen).",
                               select_mode='single',
                               pre_select=False,
                               allow_none=True,
                               )[0]
        if used_field == '':
            used_field = '{priority}'

    logger.info(f"Label docx being used for input: '{label_docx}")


    # set object type so field can accept an object - a list
    df['lines'] = ''
    df['lines'] = df['lines'].astype('object')

    for index, row in df.iterrows():
        # Substitute into the template's formatting segments rather than into one joined
        # string, so every piece of a line keeps the font, size, and weight it will really
        # be rendered in (a line mixes them - see substitute_segments()).
        # TODO what ot do if nan?  takes up only one space but might be much larger when filled in later.
        values = {fld: safe_str(df.at[index, fld]) for fld in keys}

        # fill a field in the df with the list of substituted lines (each a segment list)
        df.at[index, 'lines'] = [substitute_segments(line['segments'], values)
                                 for line in label_lines]

    # Collect max line info for alert display
    max_line_results = []
    total_rows = len(df)

    # Below is True if 'nan' found in any cell in 'keys' columns
    nan_found_in_all_rows = (df[keys] == 'nan').any(axis="columns").any(axis="rows")
    # Below creates a boolean series/mask: (df[keys] == 'nan').any(axis="columns")
    # df[keys] selects only columns in list 'keys'
    # below puts it together: creates df with 'keys' columns and rows containing 'nan' in any column:
    #   df[keys][(df[keys] == 'nan').any( axis="columns")]
    df_nan = df[keys][(df[keys] == 'nan').any(axis="columns")]

    if nan_found_in_all_rows:
        print("\n20 rows containing 'nan' somewhere in a ALL row 'key' columns")
        with pd.option_context('display.max_rows', 20, 'display.max_columns', None, ):
            print(df_nan)
        pyautobek.alert(f"'Nan' (empty) found in {len(df_nan)} rows in at least one key of '{keys}' to be substituted "
                        f"in ALL rows so analysis of all rows skipped.\n\nSee log for print of 20 rows.",
                        "'nan' found somewhere in ALL row keys")

        # there are some keys with nan so can't run on all.  Instead, find those rows with all keys non-nan.
        df_non_nan = df[(df[keys].notnull()).all(axis="columns")]  # note all columns taken
        if len(df_non_nan) != 0:
            print("ALL ROWS WITH NO NAN VALUES")
            print(df_non_nan[keys])
            print()
            print("MAX LINES IN SUBSTITUTED SCRIPT INFO FOR ALL ROWS WITH NO NAN VALUES")
            result, over = print_max_line_info(df_non_nan, 'lines', line_meta=label_lines,
                                               usable_width_pt=usable_width_pt)
            max_line_results.append(("ALL ROWS WITH NO NAN VALUES", result, len(df_non_nan), over))
        print()
        a = 1
    else:
        print("MAX LINES IN SUBSTITUTED SCRIPT INFO FOR ALL ROWS")
        result, over = print_max_line_info(df, 'lines', line_meta=label_lines,
                                           usable_width_pt=usable_width_pt)
        max_line_results.append(("ALL ROWS", result, len(df), over))
        print()

    # A row is "used" when its used_field is genuinely non-blank: not empty,
    # not whitespace, and not a missing value (read as <NA> or the literal 'nan').
    used_vals = df[used_field].fillna('').astype(str).str.strip()
    used_mask = (used_vals != '') & (used_vals.str.lower() != 'nan')

    # see explanation of expression above
    # Check for bad row where NAN is found somewhere in a key field.
    nan_found_in_select_rows = (df[keys].loc[used_mask] == 'nan') \
        .any(axis="columns").any(axis="rows")
    df_nan = df[keys][used_mask & (df[keys] == 'nan').any(axis="columns")]
    if nan_found_in_select_rows:
        print(f"\nThe following rows with non-blank '{used_field}' contain 'nan' somewhere in ALL 'key' columns")
        with pd.option_context('display.max_rows', None, 'display.max_columns', None, ):
            print(df_nan)

        print(f"\nSTATS NOT PRODUCED - USED ROWS HAVE BAD 'nan' DATA")

        pyautobek.alert(
            f"'Nan' (blank) found in {len(df_nan)} rows at least one key with non-blank '{used_field}' to be substituted so analysis "
            f"of all rows skipped.\n\n"
            "See log for list for rows.",
            "'nan' found in USED row keys")
    else:
        print(f"MAX LINES IN SUBSTITUTED SCRIPT INFO FOR USED SUB ROWS ('{used_field}' not blank)")
        df_used_counties = df.loc[used_mask]
        result, over = print_max_line_info(df_used_counties, 'lines', line_meta=label_lines,
                                           usable_width_pt=usable_width_pt)
        max_line_results.append((f"'USED' COUNTIES ('{used_field}' not blank)", result,
                                 len(df_used_counties), over))
        print()

    print("RAW LABEL DATA:")
    print(label_text)

    # Build and display alert with max line lengths and label text
    any_overflow = any(section_over for _, _, _, section_over in max_line_results)

    alert_lines = []
    alert_lines.append("CHECK LABEL LINE FITS")
    alert_lines.append("")
    alert_lines.append("--- LABEL TEMPLATE TEXT ---")
    alert_lines.append(label_text)
    alert_lines.append("")
    alert_lines.append("")
    if any_overflow:
        # Loud, and above the detail, so a line that will wrap cannot be scrolled past.
        alert_lines.append("*** ERRORS WERE ENCOUNTERED BELOW - SOME LINES DO NOT FIT ***")
        alert_lines.append("")
        alert_lines.append("")
    for section_name, section_result, row_count, _ in max_line_results:
        alert_lines.append(f"--- {section_name} ---")
        alert_lines.append(section_result)
        alert_lines.append(f"({row_count} rows of {total_rows})")
        alert_lines.append("")
        alert_lines.append("")

    alert_lines.append(f"'Used' counties based on field {used_field}")
    alert_lines.append("")
    alert_lines.append(f"Label File: {label_docx}")
    alert_lines.append("")
    alert_lines.append(f"BOE Sheet: {boe_xls}")
    alert_lines.append("")

    if usable_width_pt:
        alert_lines.append(
            f"Fit check: each line is measured piece by piece, in the real font, size, and\n"
            f"bold of every run on it (a symbol and the words beside it differ), and the\n"
            f"total compared to the label cell's usable width of {usable_width_pt:.1f} pt "
            f"({usable_width_pt / 72:.2f} in).\n"
            f"A line is called TOO WIDE past {FIT_TOLERANCE:.0%} of that width, because Word and this\n"
            f"measurement do not agree to the last fraction of a point. Suggested sizes are\n"
            f"sized to {SUGGEST_TOLERANCE:.0%} so they land with a little room to spare; shorten the text\n"
            f"instead if the suggestion is too small to read.")
    else:
        alert_lines.append(
            "Fit check: the label cell width could not be read, so line widths were\n"
            "not measured. Character counts are shown for reference only.")

    alert_message = '\n'.join(alert_lines)
    # pyautobek.alert(alert_message, "MAX LABEL LINE LENGTHS")
    scroll_box(alert_message, title="MAX LABEL LINE LENGTHS", wrap_lines=False, width=1200, height=600)


def find_keys_in_text(label_text: str) -> set:
    """Finds all {KEY} replacement tokens in a plain text string.

    Args:
        label_text (str): Text content from a label template, typically from get_label_text().

    Returns:
        set: Set of key strings found (e.g. {'{county}', '{phone}'}).
    """

    # Same pattern substitute_segments() matches with, so the keys looked up in the BOE
    # sheet and the tokens actually replaced can never drift apart.
    keys = set(_KEY_RE.findall(safe_str(label_text)))
    return keys


if __name__ == '__main__':
    INITIAL_ATTACHMENT_DIR = Path("~/Dropbox/Postcard Files/Attachments/Campaigns").expanduser()
    # LENGTH_CHECK_FIELDS_SELECT = '{priority}'
    LENGTH_CHECK_FIELDS_SELECT = '{use}'

    # Interactive GUI call (file picker + field picker dialogs). Use Run (not Debug):
    # Qt modal dialogs freeze under PyCharm's debugger on Python 3.12.
    max_label_lengths(used_field=None, initial_attachment_dir=INITIAL_ATTACHMENT_DIR, )

    # DEBUG (dialog-free) alternative — hard-code inputs, skips select_file / list_pick:
    # max_label_lengths(
    #     used_field='{use}',
    #     label_docx=Path("~/Dropbox/Postcard Files/Attachments/Campaigns/GA General 6-2026/Input/"
    #                     "LABELS-30 per page GA General 6-15-26.docx").expanduser(),
    # )

    a = 1
    # Label text is copied here.  Lines are trimmed.
    #     LABEL_TEXT = """
    # {cntytoprint} Registrar
    # Phone Num: {specialphone}
    # or {specialurl}
    # Early Voting Now thru Feb 18
    #     """

    # label_text_file = Path("~/Dropbox/Postcard Files/Attachments/Campaigns/TEST NAN3/Input/label.txt").expanduser()
    # with open(label_text_file, "r") as f:
    #     LABEL_TEXT = f.read()

    # LENGTH_CHECK_FIELDS_SELECT = '{priority}'
    # LENGTH_CHECK_FIELDS = [['county', 'phone'], 'county', 'phone']
    # LENGTH_CHECK_FIELDS = ['{cntytoprint}', '{specialphone}', '{specialurl}']
    # test = ['+'.join(x) for x in LENGTH_CHECK_FIELDS]
    # test = ['county', 'phone']

    # data = [['Small', '555-555-5555', 'x'], ['Medium', '555-555-5552 x2', 'x'], ['VeryVeryLonnnng', '555-555-5553 ext 3', '']]
    # df = pd.DataFrame(data, columns=['{county}', '{phone}', '{PRIORITY}'])

    # max_label_lengths(initial_attachment_dir=INITIAL_ATTACHMENT_DIR,)
    # label_docx="~/Dropbox/Postcard Files/Attachments/Campaigns/TEST NAN3/Input/TEST Special GOTV "
    #            "LABELS-30per page 1-3-2023.docx")
    # a = 1

    # boe_xls = "~/Dropbox/Postcard Files/Attachments/Campaigns/TEST NAN3/Input/BOE Info VA (1).xlsx",

    # df_to_check_length = pd.DataFrame
    #
    # fieldnames = []
    # for item in LENGTH_CHECK_FIELDS:
    #     if isinstance(item, list):
    #         fieldname = '+'.join(item)
    #         df_to_check_length[fieldname] = df[item].astype(str).apply(lambda x: ''.join(x), axis=1)
    #     elif isinstance(item, str):
    #         df_to_check_length[item] = df[item]
    #     else:
    #         print('ERROR - field not on dataframe')
    #     fieldnames.append(fieldname)
    #
    # df_to_check_length[LENGTH_CHECK_FIELDS_SELECT] = df_to_check_length[LENGTH_CHECK_FIELDS_SELECT]
    #
    #
    # for field in LENGTH_CHECK_FIELDS:
    #     # max_field = max(df[f"{{{field}}}"], key=len)
    #     max_field = max(df[f"{field}"], key=len)
    #     print(f"Max of '{field}' column: '{max_field}', length is {len(max_field)}")
    #
    # a=1

    # def fill_lines(df, field_list, used_field):
    #
    #     # set object type so field can accept an object - a list
    #     df['lines'] = ''
    #     df['lines'] = df['lines'].astype('object')
    #
    #     for fld in field_list:
    #         locals()[fld] = fld
    #         # newfield = df[fld]
    #
    #     for index, row in df.iterrows():
    #         # county = row['county']
    #         # phone = row['phone']
    #         # lines.append(f"{county} is the county name.")
    #
    #         df.at[index, 'lines'] = \
    #             [
    #                 f"{county} is the county name.",
    #                 f"{phone} is the phone.",
    #                 f"Together we have {county} and {phone}.",
    #                 f"Again we have {county} and {phone}."
    #             ]
    #
    #     entries = len(df.at[0, 'lines'])
    #
    #     for entry in range(entries):
    #         line_data = [x[entry] for x in df['lines'] ]
    #         max_line = max(line_data, key=len)
    #         max_length = len(max_line)
    #         print(f"line_data: {line_data}")
    #         print(f"max_line: '{max_line}', len:{max_length}\n")
    #
    #
    #     # for entry in range(entries):
    #     #     # max_field = max(df[f"{{{field}}}"], key=len)
    #     #     max_field = max(df['line'][entry], key=len)
    #     #     print(f"Max of '{field}' column: '{max_field}', length is {len(max_field)}")
    #
    a = 1
